from __future__ import annotations

import os
import re
import shutil
from pathlib import Path
from datetime import date
from typing import Dict, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).parent / "output"
ROOT = OUT / "开江中学实验学校红专八大系统建设文件包"
SCHOOL = "开江中学实验学校"
DATE_TEXT = "二〇二六年七月"
RED = "9E2A2B"
NAVY = "18324A"
GOLD = "B28A42"
LIGHT_RED = "F8EEEE"
LIGHT_BLUE = "EEF3F8"
LIGHT_GOLD = "F7F2E7"
GRAY = "666666"


def safe_name(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "_", name)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:" + edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn("w:" + key), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_run_font(run, name="Microsoft YaHei", size=Pt(12), bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_doc(doc: Document, title: str):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.55
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for style_name, size, color in [("Title", 24, RED), ("Heading 1", 16, RED), ("Heading 2", 14, NAVY), ("Heading 3", 12.5, GOLD)]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("开江中学实验学校“红专”学校发展规划 · 八大系统建设文件")
    set_run_font(r, size=Pt(9), color=GRAY)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)

    cp = doc.core_properties
    cp.title = title
    cp.subject = "开江中学实验学校红专学校发展体系"
    cp.author = SCHOOL
    cp.keywords = "三层八系统; 红专; 学校发展规划"


def add_cover(doc: Document, title: str, subtitle: str, label: str = "校级审议稿"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    r = p.add_run("开江中学实验学校")
    set_run_font(r, size=Pt(16), bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run(title)
    set_run_font(r, size=Pt(24), bold=True, color=RED)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(subtitle)
        set_run_font(r, size=Pt(14), bold=True, color=NAVY)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(7.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GOLD)
    set_cell_border(cell, top={"val": "single", "sz": 10, "color": GOLD}, bottom={"val": "single", "sz": 10, "color": GOLD}, left={"val": "single", "sz": 10, "color": GOLD}, right={"val": "single", "sz": 10, "color": GOLD})
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cp.add_run(label)
    set_run_font(rr, size=Pt(12), bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(95)
    r = p.add_run(SCHOOL + "\n" + DATE_TEXT)
    set_run_font(r, size=Pt(12), color=GRAY)
    doc.add_page_break()


def add_position_box(doc: Document, system_name: str, position: str, purpose: str):
    t = doc.add_table(rows=3, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    labels = [("系统归属", system_name), ("核心定位", position), ("根本目的", purpose)]
    for i, (a, b) in enumerate(labels):
        t.cell(i, 0).text = a
        t.cell(i, 1).text = b
        set_cell_shading(t.cell(i, 0), RED if i == 0 else NAVY)
        for r in t.cell(i, 0).paragraphs[0].runs:
            set_run_font(r, size=Pt(10.5), bold=True, color="FFFFFF")
        for r in t.cell(i, 1).paragraphs[0].runs:
            set_run_font(r, size=Pt(10.5), color=NAVY)
        t.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        t.cell(i, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_bullets(doc: Document, items: List[str], level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.65 + 0.6 * level)
        r = p.add_run(item)
        set_run_font(r, name="SimSun", size=Pt(11.5), color="222222")


def add_numbered(doc: Document, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.65)
        r = p.add_run(item)
        set_run_font(r, name="SimSun", size=Pt(11.5), color="222222")


def add_table(doc: Document, headers: List[str], rows: List[List[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=Pt(10), bold=True, color="FFFFFF")
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cells[j].text = str(text)
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i % 2 == 0:
                set_cell_shading(cells[j], "F7F9FB")
            for p in cells[j].paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, name="SimSun", size=Pt(9.5), color="222222")
    if widths:
        for row in table.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, text: str, fill=LIGHT_RED):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_shading(c, fill)
    set_cell_border(c, left={"val": "single", "sz": 18, "color": RED})
    p = c.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(title + "：")
    set_run_font(r, size=Pt(11), bold=True, color=RED)
    r = p.add_run(text)
    set_run_font(r, name="SimSun", size=Pt(11), color="333333")
    doc.add_paragraph()


def add_section(doc: Document, heading: str, paragraphs: List[str] = None, bullets: List[str] = None, level=1):
    doc.add_heading(heading, level=level)
    for text in paragraphs or []:
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = False
    if bullets:
        add_bullets(doc, bullets)


def save_doc(doc: Document, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


SYSTEMS: Dict[str, Dict] = {
    "01_红专学生成长系统": {
        "name": "红专学生成长系统",
        "position": "直接组织学生全面成长、个性发展和教育转化的核心育人运行系统",
        "purpose": "让每一名学生有成长目标、有发展路径、有真实记录、有精准支持。",
        "goals": [
            "以“三红三专”为校本育人目标，建立七至九年级递进式学生成长标准与三年成长地图。",
            "形成德育、学业、身心、能力、特长相互贯通的学生发展支持结构。",
            "建立发展性评价、分类激励、班级共建和重点学生精准关爱机制。",
            "降低重复积分、重复台账和标签化评价风险，使评价服务成长而非制造排名。",
        ],
        "tasks": [
            "学生发展标准建设：把“三红三专”转化为可理解、可观察、可发展的年级表现标准。",
            "成长课程与活动建设：统筹习惯养成、理想信念、心理成长、生涯启蒙、劳动实践、艺体科创和学生自治。",
            "班级育人建设：推动班级从纪律管理单元转向学生共同成长共同体。",
            "发展性评价与成长激励：统一综合素养评价、红专十星、红专标兵和标杆班级的标准与证据。",
            "重点学生精准关爱：建立发现、研判、建档、帮扶、转介、跟踪、复盘闭环。",
            "学生主体性建设：完善学生议事、社团自治、志愿服务和同伴互助机制。",
        ],
        "projects": ["红专学子“三红三专”培育工程", "红专十星成长激励工程", "红专十类标杆班级建设工程", "七年级六习惯筑基工程", "重点青少年“一生一策”关爱工程", "学生自治与领袖素养工程"],
        "roles": [["牵头领导", "分管德育校级领导", "统筹系统目标、重大项目和跨系统协调"], ["主责部门", "德育处（或学校核定部门）", "学生标准、班级建设、评价激励和重点学生工作"], ["执行单元", "年级、班级、团委、心理教师、导师团队", "课程活动实施、学生观察、支持转介和成长记录"], ["协同系统", "课程教学、教师发展、家校社、治理、质量、数智、保障", "提供课程、师资、家庭、程序、评价、数据和专业保障"]],
        "indicators": [["学生发展", "三红三专六维成长表现、年级增值变化", "学生自评、教师观察、代表性成果、抽样访谈"], ["班级建设", "基础达标率、特色建设质量、学生参与度", "班级发展记录、学生展示、抽样观察"], ["重点关爱", "建档完整率、措施落实率、风险变化、返校与稳定情况", "一生一策记录、会商纪要、转介与回访"], ["减负合规", "重复台账数量、教师填报时间、公开评价风险", "负担调查、文件审核、申诉记录"]],
        "support": [
            ("02_红专学子“三红三专”成长培育方案", "方案", ["六维目标释义与年级递进标准", "七年级适应筑基、八年级稳定发展、九年级励志成长", "课程课堂、活动实践、管理服务、文化环境、家校社会五域实施", "学生年度成长目标与导师支持"]),
            ("03_红专学生与班级发展性评价总则", "总则", ["一套标准、两类载体、一次采集、多方使用", "个人成长评价与班级发展评价关系", "证据类型、结果边界、隐私保护、申诉复核", "不以单次成绩和材料数量替代成长"]),
            ("04_“红专十星”学生成长评价与激励办法", "办法", ["十类成长赛道与申报条件", "过程观察、代表性成果、学生自评、师生评议", "班级初审、年级复核、学校认定", "红专标兵作为综合性高阶荣誉而非平行积分体系"]),
            ("05_“红专十类标杆班级”创建认定办法", "办法", ["五项共同基础标准", "十类特色方向与一班一主项", "创建—达标—示范三级认定", "自主申报、学生展示、动态复核、材料减量"]),
            ("06_重点青少年“一生一策”关爱服务实施方案", "方案", ["重点对象识别与数据核定", "一人一档、一生一策、导师团队、家校社联动", "长期离校和在籍不在校学生跟踪", "月度研判、专业转介、效果评估和隐私权限"]),
            ("07_学生成长档案与评价工具包", "工具包", ["学生成长画像表", "个人年度成长目标卡", "班级特色创建任务单", "一生一策个案记录表", "评价证据清单与申诉复核单"]),
        ],
    },
    "02_红专课程教学系统": {
        "name": "红专课程教学系统",
        "position": "把育人目标落实到课程、课堂、作业、教研和学习支持的核心业务系统",
        "purpose": "让课程更完整、课堂更有效、作业更适量、学习支持更精准。",
        "goals": ["建立国家课程高质量实施、校本课程特色发展和综合实践协同育人的课程结构。", "推动备课、课堂、作业、评价、辅导形成学习闭环。", "建设基于学情的分层支持和初高中衔接机制。", "审慎推进人工智能支持备课、教学和自主学习，先试点、后评估、再推广。"],
        "tasks": ["课程图谱建设：形成国家课程、地方课程、校本课程和活动课程的整体图谱。", "课堂教学改进：明确目标、任务、参与、反馈和迁移五个基本环节。", "集体备课优化：一次会议解决学情、目标、任务、作业和资源问题。", "作业治理：加强年级总量统筹、分层设计、批改反馈和跨学科协调。", "学习支持：建立卓越发展、基础巩固、学困帮扶和特殊需要支持模块。", "教研与成果：把课堂问题转化为微课题、案例、课程和共享资源。"],
        "projects": ["红专课程图谱建设工程", "集体备课与作业优化工程", "红专智学堂课堂提质试点", "红专成长领航校本课程群", "分层学习支持工程", "跨学科PBL与STEM项目"],
        "roles": [["牵头领导", "分管教学校级领导", "统筹课程、课堂、教研和质量改进"], ["主责部门", "教务处、教科室（以学校核定为准）", "课程计划、教学规范、教研组织和学业支持"], ["执行单元", "年级、教研组、备课组、学科教师", "课程实施、集体备课、课堂改进、作业协调"], ["协同系统", "学生成长、教师发展、质量、数智、保障", "提供育人目标、教师支持、评价反馈、技术和资源"]],
        "indicators": [["课程实施", "国家课程开齐开足、校本课程质量、五育融合度", "课程表、课程方案、抽样观察、学生反馈"], ["课堂质量", "目标清晰、学生参与、即时反馈、学习增值", "听评课、学生作品、学情数据、教师反思"], ["作业质量", "总量适度、分层合理、反馈及时", "作业样本、学生调查、年级协调记录"], ["教研效能", "集体备课解决问题比例、资源复用、成果转化", "一页记录单、资源包、课例和案例"]],
        "support": [
            ("02_红专课程体系建设与实施方案", "方案", ["课程理念与课程目标", "国家课程、校本课程、活动实践课程结构", "五育课程图谱与三年课程进阶", "课程准入、实施、评价和退出"]),
            ("03_集体备课与作业优化实施办法", "办法", ["四问式集体备课流程", "主备、共研、实施、复盘", "一页记录和一个共享资源包", "年级作业总量统筹与分层设计"]),
            ("04_课堂教学基本规范与学习支持指引", "指引", ["课前学情与目标", "课中任务、参与、反馈、迁移", "课后作业与辅导", "不同学科和不同学生保留合理差异"]),
            ("05_红专成长领航校本课程群实施方案", "方案", ["七年级适应与习惯", "八年级青春成长与学业发展", "九年级学业与生涯规划", "卓越发展模块与基础巩固支持模块", "书记校长领航课、班主任成长课、学科方法课、专业专题课团队授课"]),
            ("06_红专智学堂与人工智能课堂提质试点方案", "方案", ["小范围学科与班级试点", "AI辅助备课、学情诊断、分层任务和反馈", "教师最终责任、学生独立思考和数据安全", "基线—试点—中期复盘—扩展或暂停"]),
            ("07_教学准备与质量改进工具包", "工具包", ["开学前教学准备清单", "集体备课要点记录单", "年级作业总量协调表", "课堂观察简表", "学情会商和学习支持记录单"]),
        ],
    },
    "03_红专教师发展系统": {
        "name": "红专教师发展系统",
        "position": "建设教师专业能力、育人能力、协作能力和职业福祉的关键能力系统",
        "purpose": "让教师成长有方向、发展有支持、专业有舞台、工作有尊严。",
        "goals": ["建立师德、育人、教学、研究、数字素养和协同能力一体化教师发展标准。", "把书记提出的“四维四级”转化为支持型教师梯队发展工程，而非单纯等级管理。", "形成青年教师筑基、骨干教师提质、名优教师引领、领航教师辐射的梯队。", "建立教师减负、关怀、参与决策和专业意见回应机制。"],
        "tasks": ["教师发展标准：明确不同发展阶段的核心能力和支持需求。", "四维四级工程：目标、培养、评价、关爱四维协同，新秀、骨干、名优、领航四级进阶。", "校本研修：围绕学生、课程、课堂和班级真实问题开展学习。", "班主任发展：强化班级建设、学生识别、家校沟通和危机应对。", "成果孵化：支持课程、课例、案例、课题、工具和工作室建设。", "教师福祉：减少重复事务，完善关怀、荣誉、诉求和权益保障。"],
        "projects": ["红专教师“四维四级”梯队发展工程", "青蓝筑基工程", "中坚提质工程", "名优特色突破工程", "领航品牌辐射工程", "班主任专业发展工程", "教师减负关怀工程"],
        "roles": [["牵头领导", "分管教师队伍校级领导", "统筹教师发展战略、资源和评价边界"], ["主责部门", "教师发展中心、教务处、教科室、人事和工会（以学校核定为准）", "培养、研修、评价、关怀和成果支持"], ["执行单元", "年级、教研组、备课组、工作室、师徒共同体", "岗位学习、团队研修、实践改进"], ["协同系统", "课程教学、学生成长、治理、质量、数智、保障", "提供实践场景、制度、证据、技术和福祉支持"]],
        "indicators": [["梯队结构", "四级教师分布合理、发展通道清晰", "教师发展档案、年度发展对话"], ["专业能力", "课堂、育人、教研、协同和数字能力增值", "课例、案例、学生反馈、同行评议"], ["团队贡献", "带教、资源共享、团队改进和成果辐射", "师徒记录、工作室成果、共同项目"], ["教师福祉", "工作负担、职业认同、意见回应和支持获得感", "教师调查、负担审核、工会反馈"]],
        "support": [
            ("02_红专教师“四维四级”梯队发展方案", "方案", ["目标、培养、评价、关爱四维", "新秀、骨干、名优、领航四级定位", "以专业标准和发展证据认定，不以教龄自动授予名优层级", "评价结果优先用于支持、培养和资源配置，涉及绩效职称须履行专门程序"]),
            ("03_青年教师青蓝工程实施方案", "方案", ["师徒双向选择与责任", "备课、课堂、作业、育人和家校沟通五项筑基", "学期目标、月度支持、课例展示和年度复盘", "避免以材料数量评价师徒成效"]),
            ("04_班主任专业发展与支持方案", "方案", ["班级建设、学生识别、家校协同、心理安全、危机处置五项能力", "新班主任岗前培训与导师制", "班主任案例研讨和减负支持", "不把学生心理困难和风险事件简单转化为班主任扣分"]),
            ("05_教师校本研修与成果孵化办法", "办法", ["问题即课题、课堂即现场、改进即研究", "教研组、备课组、工作室分层研修", "课程、课例、案例、工具、论文和课题成果认定", "成果服务校内改进优先于外部获奖"]),
            ("06_教师减负关怀与意见回应机制", "机制", ["新增任务负担审核", "会议表格台账压减", "困难帮扶、心理支持、职业发展对话", "教师意见受理、回应、修订和反馈"]),
            ("07_教师成长评价与发展工具包", "工具包", ["教师年度发展目标卡", "四级发展自评表", "发展性面谈记录单", "师徒共同成长记录单", "成果孵化项目书", "教师负担与支持调查表"]),
        ],
    },
    "04_红专家校社协同系统": {
        "name": "红专家校社协同系统",
        "position": "联结家庭、学校、社区和专业部门共同支持学生成长的协同育人系统",
        "purpose": "形成责任互补、信息有序、资源共享、困难共解的育人共同体。",
        "goals": ["建立分年级、分主题、分对象的家长学校与家庭教育课程。", "规范家校沟通，减少情绪化、临时性和多头信息。", "建立重点学生家庭支持、属地协同和专业转介机制。", "建设社会实践、红色文化、劳动、科创、法治和职业启蒙资源库。"],
        "tasks": ["家长学校标准化运行。", "家庭教育课程与咨询支持。", "家委会依法规范参与。", "家校沟通和家长会提质。", "重点学生家庭责任与关爱支持。", "公安、卫健、民政、社区、社工等协同。", "社会实践基地和资源管理。"],
        "projects": ["红专家长学校", "分年级家庭教育课程", "重点家庭支持计划", "警校医校社校协同机制", "乡土红色与社会实践资源库", "家校沟通减负提质行动"],
        "roles": [["牵头领导", "分管德育或家校社工作校级领导", "统筹家校社政策、资源和重点个案"], ["主责部门", "德育处、家长学校办公室（以学校核定为准）", "家长课程、沟通规范、资源协同"], ["执行单元", "年级、班主任、导师、家委会、社区联络员", "日常沟通、课程实施、家庭支持"], ["协同系统", "学生成长、治理、质量、数智、保障", "个案、程序、评价、数据和专业保障"]],
        "indicators": [["家长学习", "课程覆盖、需求匹配和学习获得", "课程记录、抽样反馈、家庭实践"], ["沟通质量", "重要事项清晰度、重复通知和矛盾化解", "家长调查、投诉与回应记录"], ["重点支持", "监护责任落实、家校联系、外部协同效果", "个案记录、转介回访"], ["资源利用", "实践基地质量、安全性和学生参与", "项目任务书、风险评估、学生作品"]],
        "support": [
            ("02_红专家长学校建设实施方案", "方案", ["组织架构、课程体系、师资来源、年度安排", "七年级适应、八年级青春期、九年级升学与生涯主题", "普惠课程与重点家庭支持结合", "家长学习不以打卡次数和照片数量评价"]),
            ("03_家庭教育课程建设方案", "方案", ["家庭规则、亲子沟通、学习支持、心理健康、网络使用、法治安全", "短课、工作坊、咨询、资源包多种形式", "需求调查和课程反馈", "严禁向家长推销商业培训"]),
            ("04_家校沟通与家长会工作规范", "规范", ["信息发布归口和时间边界", "学生问题沟通的事实、影响、支持、约定四步法", "家长会减少单向成绩通报，增加成长分析和方法支持", "保护学生和家庭隐私"]),
            ("05_重点学生家校协同与监护支持办法", "办法", ["监护责任确认", "定期联系、家访和支持计划", "休学、长期请假和在籍不在校学生属地联动", "困难家庭转介与资源链接"]),
            ("06_警校医校社校协同育人机制", "机制", ["法治教育、风险预防、心理医疗、社会工作和社区支持", "常态联络与紧急联动", "个案信息最小必要共享", "重大事项会商、转介和复盘"]),
            ("07_家校社协同运行工具包", "工具包", ["家长需求调查表", "家长课程年度表", "家校沟通纪要单", "重点家庭支持记录单", "社会实践资源准入表", "协同会商单"]),
        ],
    },
    "05_红专治理系统": {
        "name": "红专治理系统",
        "position": "配置权责、统筹规划项目、规范决策制度、推动跨系统协同的组织治理系统",
        "purpose": "让目标统一、责任清楚、协同顺畅、过程可控、教师减负。",
        "goals": ["完善党组织领导的校长负责制和重大事项决策程序。", "建立三层八系统责任架构和一事一主责、多系统协同机制。", "把五年规划转化为年度项目、任务书和责任矩阵。", "建立文件、会议、表格、台账和新增任务的减负审查。", "提升学生安全风险、重大事件和整改任务的闭环治理能力。"],
        "tasks": ["规划治理：五年规划、年度计划、项目组合和资源统筹。", "责任治理：系统负责人、主责部门、协同主体、执行单元和质量归口。", "制度治理：审议、合法合规、发布、解释、修订和废止。", "项目治理：立项、任务书、进度、协调、风险、复盘和退出。", "参与治理：教师、学生、家长意见渠道。", "风险治理：重大涉生事件、警示整改、安全稳定和舆情协同。", "减负治理：清理重复事项、数据和材料。"],
        "projects": ["三层八系统协同运行工程", "年度重点项目组合管理", "教师负担清理行动", "制度文件版本治理", "重点青少年月度联合研判", "警示函整改闭环与复查", "重大涉生事件复盘机制"],
        "roles": [["领导主体", "学校党委、校级领导班子", "方向决策、重大事项、资源和问责"], ["主责部门", "党政办公室及学校核定治理部门", "规划、项目、文件、会议和责任矩阵"], ["执行单元", "处室、年级、教研组、项目组", "项目实施、过程协调和问题报告"], ["质量归口", "红专质量体系", "指标证据、评估复盘和改进建议"]],
        "indicators": [["责任清晰", "项目主责唯一、协同边界明确", "责任矩阵、争议事项记录"], ["项目效能", "节点完成、问题解决、资源匹配、负担适度", "任务书、中期协调、复盘报告"], ["制度质量", "冲突减少、版本统一、执行可理解", "现行文件清单、修订记录"], ["治理减负", "会议表格台账压减、重复采集减少", "负担清单、教师调查、审核记录"]],
        "support": [
            ("02_三层八系统协同运行办法", "办法", ["三层功能、八系统边界", "一事一主责、多系统协同、质量归口", "跨系统事项发起、会商、决策、执行和复盘", "不因新增工作另设第九系统"]),
            ("03_年度重点项目立项与推进管理办法", "办法", ["年度项目遴选标准", "项目任务书、节点、资源、风险和成果", "月度协调、中期复盘、学期决策", "保留、调整、扩大、暂停、停止五类结论"]),
            ("04_责任矩阵与跨部门协同管理办法", "办法", ["批准、负责、协同、执行、知会、质量归口", "处室、年级、教研组、班主任岗位边界", "冲突事项升级协调", "同一事项不得多头布置"]),
            ("05_文件审议发布与版本管理办法", "办法", ["文件起草、术语审核、合规审核、负担审核、审议、发布", "现行、试行、待审、参考、废止五种状态", "版本号、修订记录和现行清单", "历史废案不得覆盖最新审议结论"]),
            ("06_新增工作教师负担与风险审核办法", "办法", ["新增一表原则上合并一表", "同一数据一次采集多方使用", "教师权益、学生评价、数据隐私和安全风险先审后行", "试点未验证前不刚性考核"]),
            ("07_重大涉生事件报告处置与复盘办法", "办法", ["事件分级与首报时限", "学生保护、事实核查、家校沟通、部门联动", "舆情和隐私管理", "处置后原因分析、责任改进和制度修订"]),
            ("08_治理运行工具包", "工具包", ["年度项目申报表", "项目任务书", "RACI责任矩阵", "减负与风险审核表", "文件版本登记表", "重大事项复盘报告模板"]),
        ],
    },
    "06_红专质量体系": {
        "name": "红专质量体系",
        "position": "建立标准、采集证据、开展评价研究并推动持续改进的学校学习系统",
        "purpose": "让学校知道做得怎么样、为什么，以及下一步怎样改。",
        "goals": ["建立学生、课程、教师、家校、治理、数智和保障多维质量标准。", "坚持基线先行、少而关键、证据多元、一次采集、多方使用。", "建立学期项目复盘、年度质量报告和五年规划评估。", "推动学校重点工作与课题研究、成果转化一体化。", "避免以排名、材料数量和单一分数替代教育质量。"],
        "tasks": ["标准建设：明确什么是学生成长、有效课堂、教师发展和良好治理。", "证据治理：界定数据来源、采集周期、权限和质量。", "评价监测：形成基线、过程观察、阶段评估和年度报告。", "项目复盘：回答是否实施、遇到什么困难、怎样调整。", "研究改进：把问题转化为行动研究和可复用成果。", "成果转化：课程、制度、案例、工具、报告和推广。"],
        "projects": ["2026—2030核心指标与证据框架", "学生发展质量监测", "课堂与作业质量诊断", "教师发展质量评估", "学期重点项目复盘", "年度红专教育质量报告", "课题与实践成果转化"],
        "roles": [["牵头领导", "分管质量与科研校级领导", "统筹标准、评价、研究和改进"], ["主责部门", "教科室、质量监测部门（以学校核定为准）", "指标证据、评价工具、报告和成果转化"], ["数据主体", "八系统主责部门、年级、项目组", "按最少必要原则提供可靠证据"], ["参与主体", "教师、学生、家长和外部专业人员", "提供多元反馈和专业评估"]],
        "indicators": [["证据质量", "真实性、代表性、可比性和负担适度", "数据质量检查、抽样核验"], ["评价使用", "评价结论转化为改进的比例", "整改清单、项目调整、资源变化"], ["报告质量", "年度报告完整、克制、可解释", "质量报告、审议记录"], ["成果转化", "有效实践转化课程制度工具的数量和使用效果", "成果认定、应用反馈"]],
        "support": [
            ("02_红专教育质量评价与持续改进办法", "办法", ["质量观与评价原则", "学生、课程、教师、治理和保障质量域", "基线、监测、评估、反馈、改进流程", "评价结果不直接等同个人奖惩"]),
            ("03_学生发展质量评价指引", "指引", ["三红三专六维发展", "学业增值、身心健康、学校归属和社会责任", "学生自评、教师观察、作品表现和抽样访谈", "保护隐私、避免公开排名"]),
            ("04_课程教学与教师发展质量评价指引", "指引", ["课程实施、课堂学习、作业反馈和学习支持", "教师专业能力、团队贡献和职业福祉", "课堂观察、学习成果、同行评议和发展对话", "评价用于支持改进"]),
            ("05_学期项目复盘与年度质量报告制度", "制度", ["项目中期只回答实施、困难、支持三个问题", "学期末形成保留、调整、扩大、暂停、停止结论", "年度报告呈现进展、问题、证据和下一步", "不以宣传稿替代质量报告"]),
            ("06_学校发展与课题研究一体化成果认定转化办法", "办法", ["学校问题进入课题、项目行动提供研究现场", "课程、课例、案例、制度、工具、报告成果分类", "校内应用效果优先", "知识产权、署名和推广规范"]),
            ("07_质量证据与复盘工具包", "工具包", ["指标证据卡", "基线调查模板", "项目中期复盘表", "学期项目结论表", "年度质量报告框架", "成果认定申请表"]),
        ],
    },
    "07_红专数智系统": {
        "name": "红专数智系统",
        "position": "组织数字基础、人工智能应用、数据治理和数字安全的技术赋能系统",
        "purpose": "让技术真正服务学生成长、教师发展和治理增效，同时守住安全伦理底线。",
        "goals": ["建立统一、适度、可持续的数字平台和数据标准。", "提升教师和学生人工智能与数字素养。", "推动人工智能支持备课、学习、评价和治理的小范围高价值应用。", "建立工具准入、账号权限、数据质量、隐私保护和网络安全规范。", "避免以工具使用频次评价教师和学生。"],
        "tasks": ["数字基础和平台统筹。", "教育数据标准、口径、质量和共享。", "教师数字素养和AI教学能力。", "学生AI与数字素养课程。", "AI支持因材施教和治理减负。", "平台工具准入和退出。", "网络、账号、数据、隐私和伦理安全。", "重点学生数据最小权限管理。"],
        "projects": ["红专智学堂试点", "学生AI与数字素养课程", "教师数字能力分层培训", "一表通与数据归并行动", "平台工具准入清理", "网络欺凌与数字安全教育", "重点学生数据权限治理"],
        "roles": [["牵头领导", "分管数智或信息化校级领导", "统筹技术方向、预算、安全和伦理"], ["主责部门", "信息中心及学校核定部门", "平台、账号、数据、工具和网络安全"], ["业务主体", "八系统主责部门", "提出业务需求、确认数据口径和使用责任"], ["使用主体", "教师、学生、家长", "依法依规使用并反馈实际价值和风险"]],
        "indicators": [["应用价值", "是否减少重复劳动、改善学习和决策", "前后对比、用户反馈、项目证据"], ["数据质量", "口径一致、完整准确、重复减少", "数据字典、抽样核验"], ["数字素养", "教师学生安全、伦理、判断和创造能力", "课程作品、情境任务、调查"], ["安全合规", "权限、隐私、账号、工具准入和事件处置", "审核记录、安全日志、事件复盘"]],
        "support": [
            ("02_人工智能赋能教育应用规范", "规范", ["教师最终责任、学生独立思考、结果核验", "适用与不适用场景", "生成内容标识和学术诚信", "敏感数据禁止输入公共工具"]),
            ("03_学生人工智能与数字素养课程实施方案", "方案", ["信息判断、隐私安全、算法意识、生成式AI、数字创作", "七至九年级递进", "真实任务和作品评价", "网络欺凌、沉迷和不良信息预防"]),
            ("04_教师数字素养提升方案", "方案", ["基础工具、学情分析、AI辅助设计、数据伦理四级模块", "按需培训和岗位实践", "优秀应用案例共享", "不以登录次数和工具频次考核"]),
            ("05_学校数据治理与数据质量规范", "规范", ["数据目录、责任人、口径和更新周期", "一次采集、多方使用", "数据质量核验和纠错", "数据保留、归档和销毁"]),
            ("06_数据隐私网络安全与平台工具准入规范", "规范", ["学生和教师个人信息保护", "账号权限和最小必要", "平台安全评估、合同与退出", "网络安全事件报告和处置"]),
            ("07_数智运行工具包", "工具包", ["AI工具准入评估表", "数据字典模板", "账号权限申请表", "数据共享审批单", "数字项目价值评估表", "安全事件复盘表"]),
        ],
    },
    "08_红专保障系统": {
        "name": "红专保障系统",
        "position": "保障生命安全、身心健康、设施资源和学校正常运行的基础支持系统",
        "purpose": "让安全有底线、健康有支持、资源有保障、应急有能力。",
        "goals": ["建立校园安全风险分级管控、隐患排查、整改销号和复查机制。", "完善心理健康教育、筛查、咨询、转介、危机干预和休学复学衔接。", "保障消防、交通、食品、传染病、实验场馆和活动安全。", "完善后勤、设施、资产、经费和资源配置。", "针对警示函指出的问题开展数据核定、整改闭环和效能评估。"],
        "tasks": ["校园安全责任体系和风险清单。", "消防、交通、食品、卫生、实验和活动安全。", "学生心理健康全链条支持。", "心悦成长中心专业化运行。", "应急预案、演练和突发事件处置。", "学生伤害、欺凌和危机事件联动。", "后勤、设施、资产、财务和环境保障。", "整改任务销号、复查和质量评估。"],
        "projects": ["校园安全隐患清零行动", "心悦成长中心提质工程", "心理危机支持链建设", "消防交通专项治理", "重点场所安全升级", "应急能力提升工程", "警示函整改闭环复查"],
        "roles": [["牵头领导", "分管安全后勤校级领导", "统筹安全健康、资源和应急"], ["主责部门", "安全办、总务处、心理中心及学校核定部门", "风险排查、专业支持、设施后勤"], ["执行单元", "处室、年级、班级、场馆负责人、全体岗位", "岗位检查、首报、教育和应急"], ["协同系统", "学生成长、家校社、治理、质量、数智", "个案、外部协同、责任、评估和数据支持"]],
        "indicators": [["安全闭环", "隐患发现、整改、销号、复查及时率", "检查记录、整改证据、复查报告"], ["心理支持", "中心开放、咨询转介、危机处置和跟踪连续性", "工作台账、个案抽查、学生反馈"], ["应急能力", "预案可用、人员熟悉、演练改进", "演练记录、情境测试、复盘"], ["资源保障", "设施完好、经费匹配、服务满意", "巡检、预算执行、服务反馈"]],
        "support": [
            ("02_校园安全风险分级管控与隐患整改办法", "办法", ["风险辨识、分级、责任人和检查周期", "隐患发现、交办、整改、销号、复查", "重大风险即时报告", "学生参与安全教育与能力培养"]),
            ("03_学生心理健康教育与危机支持方案", "方案", ["普及教育、筛查观察、咨询辅导、专业转介、危机干预", "重点学生一人一策与家校协同", "休学复学和长期离校跟踪", "隐私、知情和最小必要共享"]),
            ("04_心悦成长中心建设与运行方案", "方案", ["功能定位、开放时间、人员配置和服务流程", "个体咨询、团体辅导、教师家长支持", "预约、记录、转介和督导", "场地不得闲置，服务质量定期评估"]),
            ("05_消防交通食品卫生与实验场馆安全规范", "规范", ["消防设施、通道、用电和重点场所", "电动车、上下学交通和校门秩序", "食品、饮水、传染病和卫生", "实验室、体育场馆、活动和施工安全"]),
            ("06_突发事件应急与学生伤害处置办法", "办法", ["首报、救助、保护现场、家校沟通和部门联动", "欺凌、伤害、失联、心理危机等场景", "舆情与隐私管理", "事件后复盘、支持和制度改进"]),
            ("07_后勤设施资产经费资源保障办法", "办法", ["需求申报、预算、采购、验收、使用、维护和报废", "育人项目资源优先级", "场馆设施开放与安全", "公开透明和绩效评价"]),
            ("08_保障运行工具包", "工具包", ["校园风险清单", "隐患整改销号单", "心理转介和跟踪单", "心悦中心服务登记表", "应急演练复盘表", "设施资产巡检表"]),
        ],
    },
}


def build_main_scheme(folder: Path, sys: Dict):
    title = f"{SCHOOL}{sys['name']}建设方案"
    doc = Document()
    configure_doc(doc, title)
    add_cover(doc, title, "《开江中学实验学校“红专”学校发展规划》系统主文件", "校级审议稿（建议审议后试行）")
    add_position_box(doc, "红专运行体系 · " + sys["name"], sys["position"], sys["purpose"])
    add_callout(doc, "文件定位", "本方案是该系统的母文件，规定定位、目标、任务、机制、项目、指标与配套文件。专项制度不得突破本方案确定的职责边界。")

    add_section(doc, "一、建设背景", [
        "学校已正式确定以红专文化体系、红专育人体系、红专运行体系构成“三层八系统”总体架构。本系统是红专运行体系中的专业责任单元，应当把文化价值和育人目标转化为稳定、可执行、可评价的学校实践。",
        "系统建设坚持问题导向和发展导向相统一。既回应学校当前在学生安全风险、课程课堂、教师成长、家校协同、制度协同、质量证据、数字应用和保障能力等方面的现实问题，也服务2026—2030年学校高质量发展。",
        "本方案不以增加文件、会议、表格和台账为建设成效。系统成熟度主要体现在目标是否清楚、责任是否落实、专业支持是否有效、学生和教师是否真实受益，以及学校能否依据证据持续改进。",
    ])
    add_section(doc, "二、系统定位与职责边界", [sys["position"] + "。"], sys["goals"])
    add_callout(doc, "边界规则", "本系统对本专业领域承担主责，但重大事项实行“一事一主责、多系统协同、质量体系归口评价”。不得因新增项目另设与八系统并列的新系统。", LIGHT_BLUE)

    add_section(doc, "三、建设原则", bullets=[
        "育人为本：所有任务最终解释其对学生成长、教师发展和学校改进的价值。",
        "系统协同：明确主责、协同、执行和质量归口，防止多头管理。",
        "少而关键：优先建设少数关键机制和项目，不追求文件与活动数量。",
        "试点改进：先建立基线，再小范围实施，以证据决定保留、调整、扩大、暂停或停止。",
        "减负合规：同一数据一次采集、多方使用；教师权益、学生评价、数据隐私和安全事项先审后行。",
    ])

    add_section(doc, "四、2026—2030年建设目标")
    add_table(doc, ["阶段", "时间", "核心任务", "标志性结果"], [
        ["整改筑基", "2026—2027", "核定基线、明确标准、补齐突出短板、启动少数项目", "系统母文件和核心专项办法建立；突出风险得到治理"],
        ["协同提质", "2028—2029", "机制常态运行、专业能力提升、跨系统项目协同", "形成稳定流程、有效项目和可复用成果"],
        ["定型推广", "2030", "系统评估、模式提炼、制度修订、成果转化", "形成县域初中可借鉴的红专系统建设案例"],
    ], widths=[2.5, 2.8, 6.2, 5.0])

    add_section(doc, "五、核心建设任务")
    for idx, task in enumerate(sys["tasks"], 1):
        doc.add_heading(f"（{idx}）{task.split('：')[0]}", level=2)
        detail = task.split("：", 1)[1] if "：" in task else task
        doc.add_paragraph(detail + "。实施中应明确对象、标准、责任、时间、资源与证据，避免只布置任务、不提供支持。")

    add_section(doc, "六、重点建设工程")
    rows = []
    for i, p in enumerate(sys["projects"], 1):
        rows.append([str(i), p, "围绕系统核心目标形成项目任务书；明确基线、对象、责任、资源、风险和证据；学期末复盘。"])
    add_table(doc, ["序号", "重点工程", "实施要求"], rows, widths=[1.5, 5.8, 9.2])

    add_section(doc, "七、组织运行机制")
    add_table(doc, ["责任类型", "责任主体", "主要职责"], sys["roles"], widths=[2.8, 5.0, 8.7])
    add_bullets(doc, [
        "年度规划：每年依据五年规划确定不超过若干项本系统重点项目，常规工作与项目工作分别管理。",
        "任务推进：重点项目必须形成一页任务书和责任矩阵，不要求重复制定多份同质方案。",
        "跨系统协调：一般问题由主责部门协调；涉及资源、风险、权益或重大事项的，提交治理系统会商。",
        "教师与学生参与：重大制度和评价规则试行前，应吸收教师、学生或家长代表意见。",
        "版本管理：以学校发布的现行文件清单为准，历史稿、讨论稿和AI辅助草稿不得直接作为执行依据。",
    ])

    add_section(doc, "八、质量评价与证据")
    add_table(doc, ["质量领域", "核心关注", "建议证据"], sys["indicators"], widths=[3.0, 6.3, 7.2])
    add_callout(doc, "评价原则", "基线先行、少而关键、证据多元、一次采集、多方使用、结果用于改进。没有可靠基线的指标，先建立基线，不虚构提升比例。", LIGHT_GOLD)

    add_section(doc, "九、实施进度与年度循环", bullets=[
        "暑期准备：完成上位文件、核心专项办法、责任分工、负担与风险审核。",
        "九月启动：通过干部研讨和全体教师大会发布少数重点项目、支持措施和学校承诺。",
        "九至十月：完成基线、首轮培训和项目实施。",
        "十一月：开展中期协调，只解决真实困难和资源障碍，不开展全面排名。",
        "十二月至次年一月：完成证据整理和学期复盘，决定项目去留与下一阶段安排。",
    ])

    add_section(doc, "十、风险控制与学校承诺", bullets=[
        "不以系统建设为名增加重复台账、无效材料和形式化活动。",
        "不在试点未验证前立即全面考核和长期固化。",
        "不以照片、截图、平台登录次数和材料数量代替工作成效。",
        "涉及学生隐私、教师权益、重大安全和专业转介的事项，依法依规审慎处理。",
        "学校根据师生意见和质量证据，及时调整不适宜的项目、标准和工具。",
    ])

    add_section(doc, "十一、配套文件目录")
    rows = [[str(i), title, typ, "作为本系统专项实施或操作文件"] for i, (title, typ, _) in enumerate(sys["support"], 1)]
    add_table(doc, ["序号", "文件名称", "类型", "定位"], rows, widths=[1.3, 8.2, 2.3, 4.8])

    add_section(doc, "十二、附则", [
        "本方案经学校审议后试行，由本系统主责部门会同红专治理系统、红专质量体系负责解释和组织评估。",
        "本方案与学校最新审议的《红专学校发展体系统一术语标准》《2026—2030年高质量发展规划》不一致时，以最新上位文件为准。",
    ])
    save_doc(doc, folder / f"01_{title}.docx")


def support_sections(sys: Dict, doc_title: str, doc_type: str, items: List[str]):
    action_type = "规定" if doc_type in ("办法", "规范", "制度", "总则") else "组织实施"
    purpose = f"本文件是{sys['name']}的配套{doc_type}，用于{action_type}{'、'.join([x.split('：')[0] for x in items[:3]])}等事项。"
    sections = [
        ("一、文件定位", [purpose, f"本文件服从《{SCHOOL}{sys['name']}建设方案》，不另设新的上位体系或独立评价系统。"], None),
        ("二、适用范围", ["适用于学校相关处室、年级、教研组、班级、项目组及参与本项工作的教师、学生、家长和协同单位。涉及法定职责、专业资质和上级政策的，执行国家、省、市、县最新规定。"], None),
        ("三、基本原则", None, ["育人为本、尊重主体", "标准清楚、责任明确", "过程支持、发展评价", "一次采集、材料减量", "试行复盘、动态改进", "隐私保护、风险可控"]),
        ("四、核心内容与标准", None, items),
        ("五、实施流程", None, ["准备：核定对象、基线、责任、资源和风险。", "启动：开展必要说明和培训，向相关主体讲清目标、标准和边界。", "实施：依照任务分工推进，记录少量关键事实和代表性证据。", "协调：发现跨部门、专业、安全或权益问题，及时提交主责部门和治理系统。", "复盘：根据真实成效、负担和风险形成保留、调整、扩大、暂停或停止意见。"]),
        ("六、责任分工", None, [f"{sys['name']}主责部门负责本文件组织实施和日常协调。", "相关年级、教研组、班级和岗位负责具体落实，不得重复布置同一事项。", "红专治理系统负责责任、程序、文件版本和负担风险协调。", "红专质量体系负责指标、证据、评价和学期复盘。", "红专数智系统和红专保障系统分别提供数据技术与安全资源支持。"]),
        ("七、证据与材料", None, ["只保留能够说明对象变化、实施质量、问题解决和风险控制的核心证据。", "已有平台或业务数据能够满足需要的，不再要求重复填报。", "同一活动原则上只形成一份简要记录和一组代表性成果。", "涉及学生个案、心理健康、家庭困难和教师个人信息的材料实行最小权限管理。"]),
        ("八、结果使用与边界", None, ["结果首先用于诊断、支持和改进，不直接等同于个人排名和奖惩。", "涉及绩效、职称、评优、入团、升学评价等重大结果运用的，应另行履行政策审核、民主参与和学校审议程序。", "试行期间发现负担过重、证据失真或风险不可控的，应及时调整或暂停。"]),
        ("九、监督、申诉与修订", None, ["相关主体可向主责部门提出意见或申诉，主责部门应在规定时限内核查回应。", "重要争议由学校治理协调机制处理，涉及专业问题可邀请专业人员参与。", "每学期至少复盘一次；修订后统一发布新版本，旧版本停止使用。"]),
        ("十、附则", ["本文件经学校审议后试行，由主责部门负责解释。未尽事项依据学校最新上位文件和有关政策执行。"], None),
    ]
    return sections


def build_support_doc(folder: Path, sys: Dict, n: int, title: str, doc_type: str, items: List[str]):
    full_title = f"{SCHOOL}{title}"
    doc = Document()
    configure_doc(doc, full_title)
    add_cover(doc, full_title, f"{sys['name']}配套{doc_type}", "校级试行稿（建议研讨后实施）")
    add_position_box(doc, sys["name"], f"本系统配套{doc_type}", f"把系统母文件转化为可理解、可操作、可复盘的具体规则。")
    add_callout(doc, "使用说明", "本文件正文提供核心规则；具体表单、清单和流程图可作为附件使用。学校可依据机构设置核定牵头部门名称，但不得改变职责逻辑。")
    for heading, paragraphs, bullets in support_sections(sys, title, doc_type, items):
        add_section(doc, heading, paragraphs, bullets)
    if doc_type == "工具包":
        add_section(doc, "附录：工具使用总规则", bullets=[
            "每张表只解决一个明确问题，原则上控制在一页。",
            "能够由现有数据生成的内容不要求教师重复填写。",
            "个案材料与公开评价材料分开管理。",
            "工具试用一个学期后，根据使用频率、价值和负担决定保留或取消。",
        ])
        for i, item in enumerate(items, 1):
            doc.add_heading(f"工具{i}：{item}", level=2)
            add_table(doc, ["项目", "填写内容"], [["使用对象", "由学校根据具体任务确定"], ["核心信息", "目标、事实、措施、责任、时间、结果"], ["证据要求", "只记录关键事实和代表性证据"], ["复盘意见", "保留、调整、扩大、暂停或停止"]], widths=[4.0, 12.5])
    save_doc(doc, folder / f"{n:02d}_{full_title}.docx")


def build_readme():
    title = "开江中学实验学校红专八大系统建设文件包使用说明"
    doc = Document()
    configure_doc(doc, title)
    add_cover(doc, title, "八个系统主文件及核心配套文件", "整套文件包 · 校级审议与试行使用")
    add_callout(doc, "总体结构", "本文件包依据“三层八系统”正式架构编制，包括八份系统建设方案和各系统核心配套文件。系统建设方案为母文件，配套文件负责具体实施。")
    add_section(doc, "一、文件包构成", bullets=[
        "八份系统建设方案：分别明确系统定位、目标、任务、重点工程、运行机制、指标证据和配套文件。",
        "四大育人运行系统：学生成长、课程教学、教师发展、家校社协同。",
        "四大治理保障系统：治理、质量、数智、保障。",
        "每个系统设置若干核心专项方案、办法、规范、指引和工具包，避免另行建立新的并列体系。",
    ])
    add_section(doc, "二、效力与使用顺序", bullets=[
        "学校发展总纲、2026—2030年高质量发展规划、三层纲要为上位文件。",
        "系统建设方案是各系统母文件。",
        "专项办法、规范、指引和项目方案不得突破母文件职责边界。",
        "工具表单是执行载体，不具有独立制度效力。",
        "正式实施前，应完成术语、政策、合规、教师负担、学生权益和数据安全审核。",
    ])
    add_section(doc, "三、版本状态建议", bullets=[
        "系统建设方案：校级审议后试行。",
        "涉及学生评价、教师权益、重大安全和数据隐私的专项文件：先研讨、再小范围试行。",
        "法定安全制度：依据最新法律法规和上级要求核定后发布。",
        "工具包：随项目试用，学期末决定保留、修改或取消。",
    ])
    add_section(doc, "四、统一实施要求", bullets=[
        "一个事项一个主责系统，一个牵头部门，多系统协同。",
        "同一数据一次采集、多方使用；同一活动只留一套核心证据。",
        "先定少数项目，再写专项任务书；先提供支持，再开展评价。",
        "试点未经验证，不立即与个人绩效、职称、评优等刚性挂钩。",
        "涉及学校现实数据的表述，须以经核定的官方台账为准，不在文件中使用未经核实的数据。",
    ])
    add_section(doc, "五、学校现有材料的整合归属", bullets=[
        "书记提出的教师“四维四级”作为红专教师发展系统的旗舰工程，吸收其分层培养价值，调整刚性等级和自动奖惩风险。",
        "“三红三专”作为红专育人体系下的校本育人目标模型，进入红专学生成长系统具体实施。",
        "红专十星、学生综合素养、红专标兵和标杆班级统一纳入学生发展性评价与成长激励，不再多套并行。",
        "A层、C层和全员素养课程整合为红专成长领航校本课程群及不同学习支持模块。",
        "警示函整改、一生一策、欺凌预防、心理健康和安全隐患分别归入学生成长、治理、家校社和保障系统协同推进。",
    ])
    save_doc(doc, ROOT / "00_文件包使用说明.docx")


def build_index():
    title = "开江中学实验学校红专八大系统文件总目录"
    doc = Document()
    configure_doc(doc, title)
    add_cover(doc, title, "系统主文件与核心配套文件索引", "文件目录")
    rows = []
    for folder_name, sys in SYSTEMS.items():
        rows.append([folder_name[:2], sys["name"], "01", f"{SCHOOL}{sys['name']}建设方案", "系统母文件"])
        for i, (t, typ, _) in enumerate(sys["support"], 2):
            rows.append([folder_name[:2], sys["name"], f"{i:02d}", f"{SCHOOL}{t}", typ])
    add_section(doc, "一、文件总目录")
    add_table(doc, ["系统序号", "系统", "文件序号", "文件名称", "类型"], rows, widths=[1.6, 3.4, 1.8, 8.2, 2.0])
    add_section(doc, "二、归档状态建议", bullets=[
        "本批文件统一归入‘红专运行体系—八系统建设文件’目录。",
        "系统建设方案标注为‘系统母文件’；专项方案、办法和规范标注为‘待审议/试行’；工具包标注为‘项目工具’。",
        "正式发布后建立现行版本清单；修改稿、讨论稿和书记顾问参考稿分别归入历史与参考档案，不与现行文件混用。",
    ])
    save_doc(doc, ROOT / "00_八大系统文件总目录.docx")


def build_overview():
    title = "开江中学实验学校红专八大系统建设总览"
    doc = Document()
    configure_doc(doc, title)
    add_cover(doc, title, "三层八系统 · 系统定位、目标与协同总览", "领导研讨与大会简读稿")
    add_callout(doc, "总体逻辑", "文化定向，育人立本，运行赋能；四大育人运行系统直接创造学生成长价值，四大治理保障系统提供组织、质量、技术和资源支持。")
    rows = []
    for key, sys in SYSTEMS.items():
        rows.append([key[:2], sys["name"], sys["position"], sys["purpose"]])
    add_table(doc, ["序号", "系统", "核心定位", "根本目的"], rows, widths=[1.3, 3.6, 6.5, 5.7])
    add_section(doc, "二、跨系统运行规则", bullets=[
        "学生与班级评价：学生成长系统主责，质量体系制定证据规则，治理系统保障程序，数智系统支持数据。",
        "课程与课堂改革：课程教学系统主责，教师发展系统提供专业支持，质量体系评价，数智系统试点赋能。",
        "重点青少年一生一策：学生成长系统主责，家校社、治理、保障、数智和质量系统协同。",
        "教师四维四级：教师发展系统主责，课程教学提供实践场景，治理审核权益与程序，质量体系开展发展评价。",
        "安全心理与重大事件：保障系统承担专业安全主责；学生成长、家校社和治理系统分别承担教育、家庭和组织责任。",
    ])
    add_section(doc, "三、2026年秋季优先推进", bullets=[
        "班级与学生评价：红专十星、红专标兵和十类标杆班级整合实施。",
        "教学准备：教学工作方案、集体备课、作业优化、课堂规范和学习支持。",
        "教师大会：一张图、一册简读本、一张项目表、岗位职责卡和减负承诺。",
        "风险底座：重点学生一生一策、心理健康、欺凌预防、安全隐患和重大事件复盘。",
    ])
    save_doc(doc, ROOT / "00_红专八大系统建设总览.docx")


def main():
    if OUT.exists():
        shutil.rmtree(OUT)
    ROOT.mkdir(parents=True, exist_ok=True)
    build_readme()
    build_index()
    build_overview()
    for folder_name, sys in SYSTEMS.items():
        folder = ROOT / folder_name
        folder.mkdir(parents=True, exist_ok=True)
        build_main_scheme(folder, sys)
        for idx, (title, typ, items) in enumerate(sys["support"], 2):
            build_support_doc(folder, sys, idx, title, typ, items)
    # plain-text manifest for quick preview
    manifest = []
    for p in sorted(ROOT.rglob("*.docx")):
        manifest.append(str(p.relative_to(ROOT)))
    (ROOT / "文件清单.txt").write_text("\n".join(manifest), encoding="utf-8")
    archive_base = OUT / "开江中学实验学校红专八大系统建设文件包"
    shutil.make_archive(str(archive_base), "zip", OUT, ROOT.name)
    print(f"Generated {len(manifest)} DOCX files")
    print(str(archive_base) + ".zip")


if __name__ == "__main__":
    main()
