from __future__ import annotations
from pathlib import Path
from typing import Iterable, List, Sequence
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

RED='9E2A2B'; NAVY='18324A'; GOLD='B28A42'; GRAY='666666'
LIGHT_RED='F8EEEE'; LIGHT_BLUE='EEF3F8'; LIGHT_GOLD='F7F2E7'; LIGHT_GRAY='F5F6F7'
SCHOOL='开江中学实验学校'; DATE_TEXT='二〇二六年七月'

def set_font(run, name='SimSun', size=12, bold=False, color=None):
    run.font.name=name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size=Pt(size); run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def border(cell, color='D9D9D9', size='4'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),size); e.set(qn('w:color'),color); borders.append(e)

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def add_page_number(p):
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    r._r.extend([begin,instr,end])

def configure(doc:Document, title:str, header_text='开江中学实验学校“红专”学校发展规划体系'):
    sec=doc.sections[0]
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(1.9); sec.left_margin=Cm(2.4); sec.right_margin=Cm(2.2)
    sec.header_distance=Cm(.75); sec.footer_distance=Cm(.75)
    normal=doc.styles['Normal']; normal.font.name='SimSun'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'SimSun'); normal.font.size=Pt(12)
    normal.paragraph_format.line_spacing=1.5; normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.first_line_indent=Cm(.74)
    for st_name,size,color in [('Title',24,RED),('Heading 1',16,RED),('Heading 2',14,NAVY),('Heading 3',12.5,GOLD)]:
        st=doc.styles[st_name]; st.font.name='Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(10); st.paragraph_format.space_after=Pt(6); st.paragraph_format.keep_with_next=True
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(hp.add_run(header_text), 'Microsoft YaHei', 9, False, GRAY)
    add_page_number(sec.footer.paragraphs[0])
    cp=doc.core_properties; cp.title=title; cp.author=SCHOOL; cp.subject='开江中学实验学校红专学校发展体系'; cp.keywords='三层八系统;红专;学校发展;五年规划'

def cover(doc:Document,title:str,subtitle:str='',label='校级审议稿（建议审议后试行）'):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50); set_font(p.add_run(SCHOOL),'Microsoft YaHei',16,True,NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28); set_font(p.add_run(title),'Microsoft YaHei',24,True,RED)
    if subtitle:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(9); set_font(p.add_run(subtitle),'Microsoft YaHei',14,True,NAVY)
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,LIGHT_GOLD); border(c,GOLD,'8')
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(label),'Microsoft YaHei',11,True,GOLD)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(92); set_font(p.add_run(SCHOOL+'\n'+DATE_TEXT),'SimSun',12,False,GRAY)
    doc.add_page_break()

def callout(doc,title,text,fill=LIGHT_RED):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,fill); border(c,RED,'6')
    p=c.paragraphs[0]; p.paragraph_format.first_line_indent=Cm(0); set_font(p.add_run(title+'：'),'Microsoft YaHei',11,True,RED); set_font(p.add_run(text),'SimSun',11,False,'333333')
    doc.add_paragraph()

def bullets(doc,items:Iterable[str],level=0):
    for x in items:
        p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.left_indent=Cm(.65+.55*level); set_font(p.add_run(str(x)),'SimSun',11.5)

def numbered(doc,items:Iterable[str]):
    for x in items:
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.left_indent=Cm(.65); set_font(p.add_run(str(x)),'SimSun',11.5)

def paragraph(doc,text,bold_prefix=None):
    p=doc.add_paragraph();
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix),'SimSun',12,True,NAVY); set_font(p.add_run(text[len(bold_prefix):]),'SimSun',12)
    else: set_font(p.add_run(text),'SimSun',12)
    return p

def section(doc,heading,paras:Sequence[str]|None=None,bullet_items:Sequence[str]|None=None,level=1):
    doc.add_heading(heading,level=level)
    for x in paras or []: paragraph(doc,x)
    if bullet_items: bullets(doc,bullet_items)

def table(doc,headers:List[str],rows:List[List[str]],widths=None,font_size=9.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; repeat_header(t.rows[0])
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; shade(c,NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(str(h)),'Microsoft YaHei',10,True,'FFFFFF')
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=str(v); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i%2==0: shade(c,'F8FAFC')
            for p in c.paragraphs:
                p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(2)
                for r in p.runs: set_font(r,'SimSun',font_size,False,'222222')
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Cm(w)
    doc.add_paragraph(); return t

def signature(doc,dept='开江中学实验学校'):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before=Pt(18); set_font(p.add_run(dept+'\n'+DATE_TEXT),'SimSun',12)

def save(doc,path:Path):
    path.parent.mkdir(parents=True,exist_ok=True); doc.save(path)
