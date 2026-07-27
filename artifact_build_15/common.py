from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SCHOOL = "开江中学实验学校"
RED = "8C1D1D"
NAVY = "20354F"
GOLD = "B58A3C"
LIGHT_RED = "F7ECEC"
LIGHT_BLUE = "EEF3F8"
LIGHT_GOLD = "F7F1E4"
GRAY = "666666"
LIGHT_GRAY = "F2F3F5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, keep=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepNext"))
    if keep and element is None:
        element = OxmlElement("w:keepNext")
        p_pr.append(element)
    elif not keep and element is not None:
        p_pr.remove(element)


def set_paragraph_border(paragraph, color=RED, size="8", space="4", side="bottom") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_run_font(run, east_asia="Noto Serif CJK SC", ascii_font="Aptos", size=12, bold=False, color=None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, east_asia="Noto Sans CJK SC", size=9, color=GRAY)


def new_doc(title: str, subtitle: str = "", status: str = "优化完善稿（以原稿为基础）", landscape: bool = False) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Pt(24)

    for style_name, size, color in (("Title", 26, RED), ("Heading 1", 16, RED), ("Heading 2", 14, NAVY), ("Heading 3", 12.5, NAVY)):
        st = styles[style_name]
        st.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10 if style_name != "Heading 1" else 14)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{SCHOOL}  ·  党委书记办学构想与制度参考文件优化稿")
    set_run_font(run, east_asia="Noto Sans CJK SC", size=9, color=GRAY)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run(SCHOOL)
    set_run_font(r, east_asia="Noto Sans CJK SC", size=16, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(title)
    set_run_font(r, east_asia="Noto Sans CJK SC", size=26, bold=True, color=RED)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_run_font(r, east_asia="Noto Sans CJK SC", size=14, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    r = p.add_run(status)
    set_run_font(r, east_asia="Noto Sans CJK SC", size=12, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("2026年7月")
    set_run_font(r, east_asia="Noto Sans CJK SC", size=12, color=GRAY)

    doc.add_page_break()
    return doc


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    set_paragraph_border(p, color=GOLD, size="8", space="5")


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def add_p(doc: Document, text: str, bold_prefix: str | None = None, no_indent: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if no_indent:
        p.paragraph_format.first_line_indent = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc: Document, items: Iterable[str], level: int = 0, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    if level:
        style += f" {min(level + 1, 3)}"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
        r = p.add_run(str(item))
        set_run_font(r)


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(title + "：")
    set_run_font(r, east_asia="Noto Sans CJK SC", size=11.5, bold=True, color=RED)
    r = p.add_run(text)
    set_run_font(r, size=11.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None, font_size: float = 9.5, header_fill: str = NAVY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, 100, 90, 100, 90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(str(h))
        set_run_font(r, east_asia="Noto Sans CJK SC", size=font_size, bold=True, color=WHITE)
        if widths:
            cell.width = Cm(widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 90, 90, 90, 90)
            if ridx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(str(val))
            set_run_font(r, size=font_size)
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_signature_block(doc: Document, departments: str = SCHOOL, date_text: str = "2026年7月") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(departments + "\n" + date_text)
    set_run_font(r, size=11.5)


def add_form_line(doc: Document, label: str, width_chars: int = 28) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(label + "：" + "＿" * width_chars)
    set_run_font(r, size=11)


def safe_filename(name: str) -> str:
    return name.replace("/", "／").replace("\\", "／").replace(":", "：")


def save_doc(doc: Document, out_dir: Path, filename: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / safe_filename(filename)
    doc.save(path)
    return path
