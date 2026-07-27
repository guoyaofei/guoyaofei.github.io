from __future__ import annotations

import math
import shutil
import subprocess
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont

BASE = Path(__file__).resolve().parent
PACKAGE = BASE / "output" / "开江中学实验学校党委书记15份参考文件优化完善稿"
QA = BASE / "qa"
PDF_DIR = QA / "pdf"
PNG_DIR = QA / "png"
CONTACT_DIR = QA / "contact_sheets"


def run(cmd: list[str]) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, check=True, text=True, capture_output=True)


def font(size: int):
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()


def make_contact_sheet(doc_name: str, images: list[Path], output: Path) -> None:
    thumb_w = 300
    margin = 24
    label_h = 34
    columns = 4
    thumbs = []
    for image_path in images:
        with Image.open(image_path) as im:
            rgb = im.convert("RGB")
            ratio = thumb_w / rgb.width
            resized = rgb.resize((thumb_w, max(1, int(rgb.height * ratio))))
            thumbs.append(resized)
    thumb_h = max(im.height for im in thumbs)
    rows = math.ceil(len(thumbs) / columns)
    title_h = 70
    canvas = Image.new("RGB", (margin * 2 + columns * (thumb_w + margin), title_h + rows * (thumb_h + label_h + margin) + margin), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = font(24)
    label_font = font(18)
    draw.text((margin, 18), doc_name, fill="black", font=title_font)
    for i, im in enumerate(thumbs):
        row, col = divmod(i, columns)
        x = margin + col * (thumb_w + margin)
        y = title_h + row * (thumb_h + label_h + margin)
        canvas.paste(im, (x, y))
        draw.rectangle((x - 1, y - 1, x + im.width, y + im.height), outline="#777777", width=1)
        draw.text((x, y + im.height + 5), f"第 {i + 1} 页", fill="#333333", font=label_font)
    canvas.save(output, quality=90)


def image_edge_check(path: Path) -> tuple[bool, str]:
    with Image.open(path) as im:
        gray = im.convert("L")
        mask = gray.point(lambda p: 255 if p < 245 else 0)
        bbox = mask.getbbox()
        if bbox is None:
            return False, "page appears blank"
        x0, y0, x1, y1 = bbox
        pad = 5
        if x0 <= pad or y0 <= pad or x1 >= im.width - pad or y1 >= im.height - pad:
            return False, f"content touches page edge: bbox={bbox}, size={im.size}"
        return True, "ok"


def combine_contact_sheets() -> Path:
    contact_paths = sorted(CONTACT_DIR.glob("*.jpg"))
    if len(contact_paths) != 15:
        raise RuntimeError(f"Expected 15 contact sheets, got {len(contact_paths)}")
    pages = []
    for path in contact_paths:
        with Image.open(path) as im:
            pages.append(im.convert("RGB"))
    output = QA / "ALL_CONTACT_SHEETS.pdf"
    pages[0].save(output, save_all=True, append_images=pages[1:], resolution=120.0)
    for page in pages:
        page.close()
    return output


def main() -> None:
    if QA.exists():
        shutil.rmtree(QA)
    PDF_DIR.mkdir(parents=True)
    PNG_DIR.mkdir(parents=True)
    CONTACT_DIR.mkdir(parents=True)

    docs = sorted(PACKAGE.rglob("*.docx"))
    if len(docs) != 15:
        raise RuntimeError(f"Expected 15 DOCX files, got {len(docs)}")

    report = ["15份Word文件渲染与质量检查报告", ""]
    failures = []
    total_pages = 0

    for index, docx_path in enumerate(docs, 1):
        document = Document(docx_path)
        if not document.paragraphs:
            failures.append(f"{docx_path.name}: no paragraphs")

        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(docx_path)
        ], check=True, text=True, capture_output=True)
        pdf_path = PDF_DIR / f"{docx_path.stem}.pdf"
        if not pdf_path.exists() or pdf_path.stat().st_size < 1000:
            failures.append(f"{docx_path.name}: PDF conversion failed")
            continue

        pdfinfo = run(["pdfinfo", str(pdf_path)]).stdout
        page_line = next((line for line in pdfinfo.splitlines() if line.startswith("Pages:")), "")
        pages = int(page_line.split(":", 1)[1].strip())
        total_pages += pages
        if pages < 2:
            failures.append(f"{docx_path.name}: unexpectedly short PDF ({pages} page)")

        text_path = QA / f"text_{index:02d}.txt"
        subprocess.run(["pdftotext", str(pdf_path), str(text_path)], check=True)
        extracted = text_path.read_text(encoding="utf-8", errors="replace")
        if len(extracted.strip()) < 500:
            failures.append(f"{docx_path.name}: extracted text too short")
        if "�" in extracted:
            failures.append(f"{docx_path.name}: replacement glyph found in PDF text")

        doc_png_dir = PNG_DIR / f"{index:02d}"
        doc_png_dir.mkdir()
        prefix = doc_png_dir / "page"
        subprocess.run(["pdftoppm", "-png", "-r", "100", str(pdf_path), str(prefix)], check=True)
        images = sorted(doc_png_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[-1]))
        if len(images) != pages:
            failures.append(f"{docx_path.name}: page image count {len(images)} != PDF pages {pages}")
        edge_warnings = []
        for img in images:
            ok, message = image_edge_check(img)
            if not ok:
                edge_warnings.append(f"{img.name}: {message}")
        report.append(f"{index:02d}. {docx_path.name}")
        report.append(f"    页数：{pages}；段落：{len(document.paragraphs)}；表格：{len(document.tables)}；PDF文本字符：{len(extracted.strip())}")
        report.append(f"    页面边缘检查：{'通过' if not edge_warnings else '需人工关注 ' + '; '.join(edge_warnings[:3])}")

        contact_path = CONTACT_DIR / f"{index:02d}_{docx_path.stem}.jpg"
        make_contact_sheet(docx_path.name, images, contact_path)

    combined_pdf = combine_contact_sheets()
    report.append("")
    report.append(f"Word文件：{len(docs)}份；渲染总页数：{total_pages}页；联系表：{len(list(CONTACT_DIR.glob('*.jpg')))}张。")
    report.append(f"人工检查用联系表PDF：{combined_pdf.name}。")
    if failures:
        report.append("自动检查未通过事项：")
        report.extend(["- " + item for item in failures])
    else:
        report.append("自动结构、转换、页数、文本和图片检查全部通过。")
    (QA / "QA_REPORT.txt").write_text("\n".join(report), encoding="utf-8")

    if failures:
        raise RuntimeError("QA failures:\n" + "\n".join(failures))
    print("\n".join(report[-6:]))


if __name__ == "__main__":
    main()
