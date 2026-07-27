from __future__ import annotations

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parent
OUT = BASE / "output"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "开江中学实验学校党委书记15份参考文件优化完善说明.docx"
PDF_PATH = OUT / "开江中学实验学校党委书记15份参考文件优化完善说明.pdf"

SCHOOL = "开江中学实验学校"
RED = "9C1C1C"
DARK = "333333"
GOLD = "B58A3A"
LIGHT_RED = "F8EDED"
LIGHT_GOLD = "F8F2E6"
LIGHT_GRAY = "F2F3F5"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="宋体", size=11, bold=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, "宋体", 9, color="777777")


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.3)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0.74)

    for style_name, font_name, size, color, before, after in [
        ("Title", "方正小标宋简体", 24, RED, 0, 16),
        ("Heading 1", "黑体", 16, RED, 14, 8),
        ("Heading 2", "黑体", 13, DARK, 10, 5),
        ("Heading 3", "楷体", 11.5, DARK, 6, 3),
    ]:
        st = doc.styles[style_name]
        st.font.name = font_name
        st._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("开江中学实验学校党委书记15份参考文件优化完善说明")
    set_run_font(r, "宋体", 9, color="777777")
    add_page_number(sec.footer.paragraphs[0])
    return doc


def add_paragraph(doc, text, *, indent=True, bold_prefix=None, color=DARK, size=11):
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, "黑体", size, True, color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, "宋体", size, False, color)
    else:
        r = p.add_run(text)
        set_run_font(r, "宋体", size, False, color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("● ")
        set_run_font(r, "宋体", 10.5, True, RED)
        r = p.add_run(item)
        set_run_font(r, "宋体", 10.5, False, DARK)


def add_callout(doc, title, text, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "：")
    set_run_font(r, "黑体", 10.5, True, RED)
    r = p.add_run(text)
    set_run_font(r, "宋体", 10.5, False, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_simple_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], RED)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(h)
        set_run_font(r, "黑体", font_size, True, "FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "FAFAFA")
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, "宋体", font_size, False, DARK)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return table


FILES = [
    {
        "title": "1.《“红专十星”学生评选活动实施方案》",
        "retain": ["“人人有赛道、个个可出彩”的基本理念", "学习、品德、体育、艺术、文学、劳动、管理、奉献、创新、励志十大成长类别", "日常表现70分、代表性成果20分、师生评议10分的主体结构", "学生申报、班级初评、年级复核、学校审定、公示表彰的五步流程"],
        "opt": ["将“一人只能评一星”调整为“原则上一生一主项，确有突出表现可审核增报一项”，防止荣誉垄断，也避免机械限制。", "取消每日全量积分和重复台账，过程评价主要使用已有教育教学记录及不超过3项关键事实。", "明确没有校外证书的学生，也可凭持续进步、劳动服务、同伴互助和真实成长参评。", "补充申诉复核、隐私保护及一般违纪后的综合判断，避免一项问题永久否定学生。", "增加学生申报表和班级初评表，班主任拿到文件即可直接组织。"],
        "say": "十星的基本设计没有改变，主要把重积分、重材料的评选，优化为重成长、重事实、面向全体学生的激励机制。",
    },
    {
        "title": "2.《“红专十星”评选方案》",
        "retain": ["十类星级及五育融合导向", "过程表现60分、实绩成果40分的主体结构", "全员普惠、分类评价及五步评审流程"],
        "opt": ["统一修正重复条款、同一项目出现两种分值、个别类别总分超过100分等技术问题。", "每一星级统一为三个过程维度60分和一个成长成果维度40分，评分口径更清楚。", "弱化年级前10名、前50名等排名，学习之星增加学习习惯、方法、个人增值和同伴互助。", "同一成果只认定一次，不再按层级机械累加；个人明显进步可与竞赛成果等值认定。", "将成果材料控制在3项以内，增加完整的学期实施时间表。", "明确星级不得作为入团、学生干部选拔或其他重要事项的唯一依据。"],
        "say": "类别和核心理念完全保留，重点解决了分值冲突、排名依赖、材料过多和结果使用过刚的问题。",
    },
    {
        "title": "3.《“红专十大标杆班级”评选量化评分细则》",
        "retain": ["十类标杆班级设置", "学期创建周期、过程评价、量化评价和动态复核", "班级特色建设、学生参与和成果展示"],
        "opt": ["将十套重复评分项目整合为共同基础40分、特色建设40分、学生主体与展示20分。", "共同基础统一评价安全秩序、尊重包容、学生参与和基本班风学风。", "每班原则上选择一个主创建类别，避免同时申报多项、准备多套材料。", "证据压缩为“一页创建记录＋三项代表性成果＋学生抽样访谈”。", "月度观察只用于诊断和支持，不再累计排名或反复授牌。", "发生事件时重点看发现、报告、保护、处置和整改，不因个别学生问题机械否定全班。", "增加创建级、达标级、示范级三级认定和可直接使用的评分记录表。"],
        "say": "保留十大标杆和量化评价，把十套重复式检查表整合成共同基础加一项特色，既能评价，也能减轻班主任负担。",
    },
    {
        "title": "4.《“红专十大标杆班级”创建认定实施方案》",
        "retain": ["班级自主申报、学期持续培育", "学生参与创建与自主展示", "师生共同认定、授牌后动态复核"],
        "opt": ["将多项同时申报调整为“一班一主项”，确保特色建设真正做深。", "建立诊断选项、申报立项、持续建设、中期会诊、自主展示、认定改进六步流程。", "将学生评审权重由70%调整为30%，教师与牵头部门50%，年级和学校复核20%。", "把中期“检查材料”优化为会诊困难、提出不超过3条改进建议。", "明确个别学生问题不自动连带全班；只有瞒报、纵容欺凌、公开羞辱、材料造假等经查实情况才取消资格。", "增加一页式创建任务书，不再要求班级另写长篇方案。"],
        "say": "保留了学生自主创建和展示的亮点，主要对多项挂牌、学生评审权重过高和集体连带进行了稳妥调整。",
    },
    {
        "title": "5.《学生综合素养考评及红专标兵评选实施方案》",
        "retain": ["道德品行40分、学业提升35分、社会责任25分", "班级、年级、学校三级评议", "班级、年级、校级红专标兵三级荣誉"],
        "opt": ["将校纪管理与成长评价适度分开，避免同一行为被处分、扣分和取消评优多重重复处理。", "综合评价既记录问题事实，也记录教育措施、整改态度和后续变化。", "学业板块增加学习态度、方法、个人增值和代表性成果，降低排名依赖。", "不再公开详细扣分和敏感行为，评价结果定向反馈学生本人和监护人。", "明确红专标兵是综合性高阶荣誉，不与十星形成另一套日常积分。", "补充评选名额建议、完整申诉程序和学期综合素养反馈表。"],
        "say": "40＋35＋25和红专标兵主体结构不变，主要让评价既守住行为底线，也能看见学生整改和进步。",
    },
    {
        "title": "6.《学生综合素养考评及红专标兵评选方案政策对标审查与优化报告》",
        "retain": ["政策对标、逐项审查、优势分析、修改建议和审查结论的基本结构"],
        "opt": ["将文件性质调整为学校内部政策对标审查材料。", "删除无法核验的“全国资深专家终审”“98分”“国家级示范”等表述。", "由单纯证明“全部符合”改为同时呈现优势、风险和优化结论。", "增加立德树人、过程增值、反唯分数、教育惩戒、公平申诉、隐私保护、教师负担和体系衔接八个审查领域。", "增加准备、试行、复盘、修订推广四阶段安排。"],
        "say": "政策对标功能保留，但将不可核验的外部审定证明，调整为学校能够正式使用、风险更小的内部审查报告。",
    },
    {
        "title": "7.《红专教师阶梯式分层培养与层级认定实施方案》",
        "retain": ["新秀、骨干、名师、领军名师四级梯队", "首次认定参考参工年限", "后续晋级看能力和业绩、分层培养和动态发展"],
        "opt": ["把参工年限由自动定级依据调整为首次分组的重要参考，同时核验履职、能力、意愿和贡献。", "名师和领军名师必须有团队带动、课程建设或成果转化事实，不能仅凭教龄。", "补充本人申报、团队评议、学校认定、发展面谈和复核程序。", "为四级分别配置年度必做任务、学校支持和代表性成果。", "晋级从课堂、育人、教研课程、团队贡献等多个领域综合判断。", "将年度简单降级优化为发展性评价、改进期和规范审议。", "明确层级主要用于培养和资源支持，不作为绩效、职称和评优的唯一依据。"],
        "say": "四级阶梯没有改变，主要是从按工龄定身份，优化为参考工龄判断发展阶段，再依据能力和贡献培养、晋级。",
    },
    {
        "title": "8.《教师“四维四级”梯队发展体系建设方案》",
        "retain": ["目标、培养、评价、关爱四维", "新秀、骨干、名优、领航四级", "青蓝筑基、中坚提质、特色突破、品牌辐射四项培养工程"],
        "opt": ["明确四级是专业发展阶段，不是行政身份、薪酬等级或固定待遇等级。", "补充年度目标确定、实践改进、中期面谈、成果展示和年度评价的完整周期。", "建立师德育人、课堂教学、教研课程、团队服务、发展增值五类评价领域。", "由重成绩、论文和奖项，转向课堂观察、学生作品、育人案例、课程资源、带教效果和团队贡献。", "增加文化课、艺体、班主任、行政兼课和教辅岗位的分类评价边界。", "评价结论形成优势、改进方向、支持建议和下一年度目标，不公开个人分数排名。", "增加教师复核、权益保护、困难帮扶和新增负担审核。", "晋级原则上每两至三年一次，自主申报，不强制人人竞争。"],
        "say": "四维四级核心构想保留最完整，重点补足怎样年度运行、怎样分类评价、怎样提供支持，以及怎样避免变成刚性等级管理。",
    },
    {
        "title": "9.《“三红三专”一体化课程体系实施方案》",
        "retain": ["三红铸魂、三专赋能两大课程集群", "国家课程、校本课程、拓展选修、社团实践四类课程形态", "七年级筑基、八年级发展、九年级圆梦三年递进", "校史、本土红色文化、学习方法、心理、劳动、艺体、科创和人工智能等内容"],
        "opt": ["明确国家课程必须开齐开足，校本课程不得挤占国家课程课时。", "为四类课程明确对象、课时、责任边界和建设要求。", "设置分阶段建设优先级，避免所有课程构想在一年内同时铺开。", "增加课程准入，一页说明对象、目标、课时、师资、场地、安全和评价。", "涉及校外机构、收费、AI平台、个人信息、实验和户外活动的，必须专项审核。", "增加单课“情境与目标—学习与实践—表达与反馈—行动与延伸”的操作结构。", "增加课程复盘和保留、调整、扩大、暂停、停止五类退出结论。"],
        "say": "大架构和特色内容都保留，重点补上先做什么、谁来做、占什么课时、怎样开课、怎样评价和效果不好怎样退出。",
    },
    {
        "title": "10.《集体备课规范化管理制度》",
        "retain": ["每周集体备课", "上周学情复盘、下周主备说课、练习作业研讨、资源共享和任务分工的主体闭环"],
        "opt": ["明确统一的是课标进度、核心目标、重难点、质量底线和作业总量，不要求课堂流程、教案、课件完全相同。", "允许教师根据不同班级学情进行二次设计。", "每次备课聚焦3—5个关键问题，不以议程和材料数量衡量质量。", "每次只形成一张记录表和一个共享资源包，取消重复打印归档。", "明确常规备课时长和小规模学科隔周、跨年级联合等灵活方式。", "补充新授、复习、讲评、实验实践、毕业年级等不同类型备课重点。", "管理由全覆盖签字检查改为抽样观察，重点看是否解决真实教学问题。", "把学生受益和教师负担同时纳入质量评价。"],
        "say": "严格要求没有降低，但从统一教案、统一课堂，优化为统一质量底线、允许教师二次设计，并大幅减少重复材料。",
    },
    {
        "title": "11.《集体备课活动记录表》",
        "retain": ["基本信息、学情复盘、下周教学研讨、作业练习、任务分工和签字栏目"],
        "opt": ["改为横向、结构化、可直接填写的表格。", "复盘只填写关键事实、共同判断、补救安排和责任人，不重复抄写教案。", "增加“允许教师依据班情调整的内容”，与二次设计制度衔接。", "将作业细分为课堂练习、基础作业、选择提升、个别支持和小测周练，并记录时长和反馈方式。", "增加任务责任人、完成时间、共享位置和下次需要反馈的问题。", "新增实施后简要反馈，使备课记录形成闭环，而不是会后即归档。"],
        "say": "原表主要栏目保留，但改成一张表完成复盘、决策、分工和后续反馈，老师不再重复抄写已有材料。",
    },
    {
        "title": "12.《初三A层次班级16周学业指导与生涯规划专题课教学提纲》",
        "retain": ["学习态度、学习能力、学习方法、学习策略四维框架", "面向学业基础较好学生的16周课程", "学业突破、心态、方法、考试策略和生涯规划主题"],
        "opt": ["明确A层是阶段性学习支持分组，不是固定身份和公开标签。", "进入与调整结合近期需要、本人意愿和教师判断，不只看一次考试排名。", "删除“满分公式”等过度承诺，调整为稳定基础、突破瓶颈和可持续卓越。", "每周增加明确课堂任务和成果，如目标卡、得分结构图、错因图谱、稳态应对卡和生涯探索表。", "文科、理科、英语专题由相应学科骨干共同承担。", "设置基线、期中、期末三次同一量表诊断。", "持续焦虑和情绪困扰及时转介心理专业支持。", "不形成公开排名，期末提供个人四维成长反馈。"],
        "say": "16周主体内容基本保留，重点降低标签化和过度应试色彩，让课程既能拔尖，也兼顾心态、生涯和持续发展。",
    },
    {
        "title": "13.《初三C层次班级16周学业指导与生涯规划专题课教学提纲》",
        "retain": ["理解和共情基础较弱学生", "态度、能力、方法、策略四维框架", "先稳心态、再养习惯、再提能力", "基础优先、低压力、小步进步和多元生涯指导"],
        "opt": ["明确C层是阶段性支持分组，不使用“差生”“后进生”等公开标签。", "保留理解学生委屈、承认学生努力的开篇思想，压缩为更适合课堂使用的核心表达。", "第一课增加匿名表达和保密约定。", "每周配置一个低门槛、可完成的成果，如极简复盘卡、情绪求助卡、基础清单和路径信息卡。", "建立普惠课堂、导师短时、家校协同、专业转介四层支持。", "对持续低落、自伤风险和严重家庭冲突增加报告保护和转介要求。", "升学和职业教育信息以当年正式政策为准，不作保证性承诺。", "期末采用进步描述，不进行公开排名和固定等级标签。"],
        "say": "最核心的共情和兜底育人思想完整保留，主要进一步保护学生尊严，并把一次课程扩展为连续支持机制。",
    },
    {
        "title": "14.《全员素养提升暨创新拔尖培育校本课程提纲》",
        "retain": ["创新潜质发展与全员基础素养两大模块", "全年32课时", "创新思维、科创实践、品格习惯、学习方法、身心健康和生涯教育", "党委书记领航课程的特色定位"],
        "opt": ["将党委书记全程亲授优化为书记领航关键主题、专业教师团队协同实施。", "创新和基础模块实行动态选择，不以固定班级或一次考试判断学生价值。", "为两个模块设置共同基础课程，保证所有学生都获得文化、规则、方法、健康、责任、数字素养和生涯教育。", "每一课均配置具体学生成果，减少单向讲座和形式化心得。", "增加AI、实验、无人机、校外实践和个人信息的安全边界。", "每学期根据学生获得感、课程负担和实际成果复盘。", "明确授课工作量和课程团队分工，保证长期可持续。"],
        "say": "书记领航、全员提升、创新拔尖的核心不变，主要优化为书记领航、专业团队共同实施，并增加动态选择和成果要求。",
    },
    {
        "title": "15.《“十五五”（2026—2030）发展规划建议》",
        "retain": ["红专文化、三红三专育人", "扩优提质、健康第一、强师兴教和数字赋能", "课程教学、学生评价、教师发展、家校社、安全心理和质量提升等主要方向"],
        "opt": ["明确文件是党委书记办学构想和正式规划编制的参考建议稿，不与学校正式规划并行。", "不直接使用未经核实的比例和精确提升数据，提出2026年先核定真实基线。", "补充学校发展基础、现实问题和2030年应达到的改变。", "形成2026—2027整改筑基、2028—2029协同提质、2030模式定型三个阶段。", "将十星、标杆班级、四维四级和分层课程纳入12项重点工程，不新增并列上位体系。", "每年从12项工程中选择6—8项重点推进，避免所有事项同时全面铺开。", "增加项目任务书、中期协调、学期复盘、年度报告和规划调整机制。", "指标先确定方向和目标类型，具体比例在基线核定后审议确定。", "增加2026至2030逐年重点安排。"],
        "say": "原有战略方向没有改变，主要把宏大的五年设想转化为有基线、有阶段、有重点工程、有年度选择和调整机制的规划建议。",
    },
]


def build_doc():
    doc = setup_document()

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SCHOOL)
    set_run_font(r, "黑体", 20, True, RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("党委书记15份参考文件优化完善说明")
    set_run_font(r, "方正小标宋简体", 26, True, RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("原稿主体保留 · 针对性完善 · 实施操作优化 · 风险边界校正")
    set_run_font(r, "楷体", 13, False, GOLD)
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("交流说明稿")
    set_run_font(r, "黑体", 13, True, DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年7月")
    set_run_font(r, "宋体", 11, False, "666666")
    doc.add_page_break()

    doc.add_heading("一、总体说明", level=1)
    add_paragraph(doc, "本轮优化以学校党委书记发来的15份办学构想、制度参考稿和内部建设材料为基础，未改变“红专十星、十大标杆班级、三红三专课程、教师四维四级、分层学业指导、全员素养课程和十五五发展”等主要办学构想。优化工作的核心不是另起炉灶，而是使原稿更加规范、稳妥、清晰和便于执行。")
    add_callout(doc, "总体处理原则", "保留主体和主要内容，只对重复冲突、操作不足、材料负担、程序权益和实施节奏进行有限度优化。", LIGHT_RED)
    add_bullets(doc, [
        "从办学构想转化为可执行流程：补充对象、责任、步骤、时间、证据、复核和表单。",
        "从多套积分和全量留痕转化为关键事实和代表性证据：避免教师、学生和家长被材料牵着走。",
        "从简单排名、扣分和一票否决转化为发展性评价和规范程序：既守底线，也看整改、进步和实际贡献。",
        "从一次全面铺开转化为试行、复盘、完善、推广：把书记构想做实，而不是急于一次定型。",
    ])

    doc.add_heading("二、15份文件优化速览", level=1)
    summary_rows = []
    for idx, item in enumerate(FILES, 1):
        short_title = item["title"].split("《", 1)[1].rstrip("》")
        summary_rows.append([str(idx), short_title, item["say"]])
    add_simple_table(doc, ["序号", "文件", "优化后的核心变化"], summary_rows, widths=[1.3, 6.2, 9.0], font_size=8.3)

    doc.add_page_break()
    doc.add_heading("三、15份文件逐项优化说明", level=1)
    for idx, item in enumerate(FILES, 1):
        if idx > 1:
            doc.add_page_break()
        doc.add_heading(item["title"], level=2)
        doc.add_heading("（一）原稿主体保留", level=3)
        add_bullets(doc, item["retain"])
        doc.add_heading("（二）本轮主要优化", level=3)
        add_bullets(doc, item["opt"])
        add_callout(doc, "与书记交流时可概括", item["say"], LIGHT_GOLD)

    doc.add_page_break()
    doc.add_heading("四、跨文件统一完成的优化", level=1)
    add_bullets(doc, [
        "统一正式封面、文件名称、标题层级、正文格式、页眉页脚和表格样式。",
        "将重要流程、责任、评分、进度和证据要求改为结构化表格，提升查阅效率。",
        "实施类文件尽量配置一页式任务书、申报表、评价表、反馈表或行动卡，便于直接使用。",
        "删除“AI生成”“可信度10分”等不宜进入正式学校文件的说明。",
        "删除文件末尾不必要的参考来源介绍，正文只保留与实施、审议和使用直接相关的内容。",
        "统一压减重复台账、重复公示、重复证明和过度照片留痕，强调一次采集、多方使用。",
        "统一补充学生和教师的陈述、申诉、复核、隐私和结果使用边界。",
        "统一将重大制度建议定位为试行、复盘、修订后再正式定型，避免一开始就形成过度刚性的管理压力。",
    ])

    doc.add_heading("五、与党委书记交流时可直接采用的总体表述", level=1)
    add_callout(doc, "建议表述", "这次优化没有改变您提出的红专十星、十大标杆班级、三红三专课程、教师四维四级、分层学业指导和十五五发展等主体构想。主要做了四方面完善：一是修正少量重复、分值冲突和前后口径不一致；二是补充责任主体、时间节点、实施流程和操作表格；三是适当降低重复积分、材料留痕和教师负担；四是完善学生教师权益、申诉复核、隐私保护和分阶段试行机制。总体上是保留您的办学思想，让这些构想更加规范、稳妥，也更便于学校教师直接执行。", LIGHT_RED)

    doc.add_heading("六、交流时需要特别说明的边界", level=1)
    add_bullets(doc, [
        "本轮优化稿是对书记原稿的完善版本，不代表所有文件应当同时正式发布。",
        "涉及学生处分、教师绩效职称、心理健康、安全、数据和经费的条款，正式发布前仍需由对应部门进行专项核定。",
        "文件中出现的时间、部门名称、人数、比例和基线数据，应结合学校实际和当年工作安排最后确认。",
        "优化的目的不是削弱管理要求，而是让要求更清楚、程序更稳妥、证据更真实、执行更可持续。",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("开江中学实验学校党委书记15份参考文件优化完善说明")
    set_run_font(r, "楷体", 10.5, False, "666666")

    doc.save(DOCX_PATH)
    return doc


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
