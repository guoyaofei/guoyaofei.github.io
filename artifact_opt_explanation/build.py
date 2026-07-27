from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = Path(__file__).resolve().parent / 'output'
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / '开江中学实验学校党委书记15份参考文件优化完善说明.docx'

RED = '8B1E1E'
NAVY = '19324D'
GOLD = 'B58A34'
LIGHT_RED = 'F7ECEA'
LIGHT_GOLD = 'F7F1E3'
LIGHT_BLUE = 'EDF3F8'
GRAY = '666666'
WHITE = 'FFFFFF'

FILES = [
    {
        'title': '《“红专十星”学生评选活动实施方案》',
        'keep': [
            '保留“人人有赛道、个个可出彩”的普惠激励理念。',
            '保留学习、品德、体育、艺术、文学、劳动、管理、奉献、创新、励志十大类别。',
            '保留“日常表现70分＋代表性成果20分＋师生评议10分”的评价结构。',
            '保留学生申报、班级初评、年级复核、学校审定、公示表彰五步流程。',
        ],
        'changes': [
            '把“一人只能评一星”优化为“原则上一生一主项，确有突出表现可审核增报一项”，防止荣誉垄断，也避免机械限制。',
            '取消每日、每项持续累计的繁重积分流水账，改用已有教育教学记录和不超过三项关键事实。',
            '明确没有校外获奖证书的学生，也可以凭持续进步、劳动服务、同伴帮助和真实成长参评。',
            '补充学生申诉、家长复核和隐私保护程序，不公开详细扣分、家庭困难及其他敏感信息。',
            '增加学生申报表和班级初评表，使班主任能够直接组织实施。',
        ],
        'summary': '没有改变“红专十星”的基本设计，而是把它由重积分、重台账的评选活动，优化为重成长、重事实、普通学生也有机会的激励机制。',
    },
    {
        'title': '《“红专十星”评选方案》',
        'keep': [
            '保留十大星级设置和全员普惠、分类评价、五育融合导向。',
            '保留“日常过程60分＋代表性成果40分”的主体结构。',
            '保留学生申报、班级初评、年级复核、分类审定和公示表彰流程。',
        ],
        'changes': [
            '统一修正原稿中重复条款、同一项目两种分值、部分星级总分超过100分等技术问题。',
            '将每一星级统一为三个过程维度共60分、一个成长成果维度40分，标准更加清楚。',
            '弱化年级前10名、前50名、前100名等排名，学习之星更多关注学习习惯、学习方法、个人增值和同伴互助。',
            '同一成果只认定一次，不再按国家、省、市、县、校机械累加；个人进步可与竞赛成果等值评价。',
            '成果材料原则上不超过三项，减少学生、家长和教师集中制作材料。',
            '增加第1周至第18周完整运行时间表，并明确星级不得作为入团、干部选拔、升学评价的唯一依据。',
        ],
        'summary': '十星的类别和核心理念完全保留，重点解决了分值冲突、过度依赖排名、材料过多和结果运用过刚的问题。',
    },
    {
        'title': '《“红专十大标杆班级”评选量化评分细则》',
        'keep': [
            '保留十类标杆班级的完整设置。',
            '保留一个学期为一个创建周期、过程评价和量化评价。',
            '保留班级特色建设、学生参与、成果展示和动态复核。',
        ],
        'changes': [
            '把十个类别中大量重复项目整合为“共同基础40分＋特色建设40分＋学生主体与展示20分”。',
            '共同基础统一评价安全秩序、尊重包容、学生参与和基本班风学风。',
            '每班原则上选择一个主创建类别，不再鼓励同时准备多套材料、争取多块牌匾。',
            '证据压缩为“一页创建记录＋三项代表性成果＋学生抽样访谈”。',
            '取消月度累计排名，月度观察只用于发现问题和提供支持。',
            '发生事件时不直接对全班一票否决，重点评价是否及时发现、报告、保护、处置和整改。',
            '增加创建级、达标级、示范级三级认定和可直接使用的评分记录表。',
        ],
        'summary': '保留十大标杆和量化评价，但把十套重复式检查表整合成一套共同基础加一项班级特色，既能评，也不会把班主任压在材料中。',
    },
    {
        'title': '《“红专十大标杆班级”创建认定实施方案》',
        'keep': [
            '保留班级自主申报、一个学期持续培育、学生参与创建、学生自主展示、师生共同认定和动态复核。',
        ],
        'changes': [
            '将“班级可同时申报多项”调整为“一班一主项”，确保每个班真正形成一项有质量的特色。',
            '建立“诊断选项—申报立项—持续建设—中期会诊—自主展示—认定改进”六步流程，并明确周次、班级任务和学校支持。',
            '将学生评审占70%调整为学生30%、教师和牵头部门50%、年级及学校复核20%。',
            '中期环节由检查材料调整为会诊困难、提出不超过三条改进建议。',
            '明确个别学生问题不自动连带否定全班；只有迟报瞒报、放任欺凌、公开羞辱、材料造假等经规范程序查实后，才取消资格。',
            '增加一页式创建任务书，班级不必再另写长篇实施方案。',
        ],
        'summary': '保留了“学生自主创建、自主展示”的亮点，主要对多项挂牌、学生评审权重过高和集体连带问题进行了稳妥调整。',
    },
    {
        'title': '《学生综合素养考评及红专标兵评选实施方案》',
        'keep': [
            '保留道德品行40分、学业提升35分、社会责任25分的基本结构。',
            '保留班级、年级、学校三级评议。',
            '保留红专标兵的班级、年级、校级三级荣誉。',
        ],
        'changes': [
            '将大量具体违纪扣分条款从综合素养评价中适度分离；校纪问题按正式制度处理，成长评价主要记录事实、教育措施、整改和后续进步。',
            '避免同一行为“处分一次、积分再扣一次、评优再否定一次”的重复惩罚。',
            '学业评价不再主要依据班级和年级排名，增加学习态度、学习方法、个人增值和代表性学习成果。',
            '不再逐级公开完整扣分明细和敏感行为，评价结果主要反馈学生本人和监护人。',
            '明确红专标兵是学校综合性高阶荣誉，不与“红专十星”形成另一套日常积分。',
            '把红专标兵条件调整为综合素养优秀或进步显著、品德可靠、学习态度端正、具有持续实践并能真实讲述成长。',
            '增加评选名额比例建议、申诉程序和学期综合素养反馈表。',
        ],
        'summary': '保留40＋35＋25和红专标兵主体结构，重点把纪律管理与成长评价适度分开，让评价既有底线，也能看见学生整改和进步。',
    },
    {
        'title': '《学生综合素养考评及红专标兵评选方案政策对标审查与优化报告》',
        'keep': [
            '保留政策对标、逐项审查、总体评价、优势分析、修改建议和审查结论。',
        ],
        'changes': [
            '将文件性质由“外部专家终审报告”调整为“学校内部政策对标审查与优化报告”。',
            '删除无法核验的“全国资深专家审定”“终审98分”“国家级示范标准”“可作为区域样板”等表述。',
            '由只写“全部符合”改为同时列出原稿优势、潜在风险和优化结论。',
            '增加立德树人与五育、过程与增值、反唯分数、教育惩戒、公平申诉、隐私保护、教师负担、体系衔接八个审查领域。',
            '增加准备期、一个年级试行、学期复盘、修订推广四阶段建议。',
        ],
        'summary': '政策对标功能保留，但把不可核验的“外部专家终审证明”改成学校可以正式使用、风险更小的内部审查报告。',
    },
    {
        'title': '《红专教师阶梯式分层培养与层级认定实施方案》',
        'keep': [
            '保留红专新秀、红专骨干、红专名师、红专领军名师四级梯队。',
            '保留首次认定参考参工年限、后续晋级看业绩和能力、分层培养与动态发展的主体思路。',
        ],
        'changes': [
            '将“完全按工龄自动定级”调整为参工年限是首次分组的重要参考，同时核验岗位履职、专业能力、个人意愿和实际贡献。',
            '名师和领军名师不再仅凭教龄自动获得，必须具有团队带动、课程建设或成果转化事实。',
            '增加教师本人申报、教研组专业评议、学校综合认定、发展面谈和复核程序。',
            '为四个层级分别配置年度必做任务、学校支持资源和代表性成果要求。',
            '晋级不再使用单一分数，而是从课堂教学、育人实践、教研课程、团队贡献等领域综合判断。',
            '把“一年一降级”“直接清零降为新秀”调整为发展性评价、改进期和规范审议。',
            '明确层级主要用于培养和资源支持，不作为绩效、职称、评优和岗位遴选的唯一依据。',
        ],
        'summary': '四级阶梯没有改变，主要把“按工龄自动定身份”优化为“参考工龄确定发展阶段，再依据能力和贡献提供培养与晋级”。',
    },
    {
        'title': '《教师“四维四级”梯队发展体系建设方案》',
        'keep': [
            '完整保留目标体系、培养体系、评价体系、关爱体系“四维”。',
            '完整保留新秀、骨干、名优、领航“四级”。',
            '保留青蓝筑基、中坚提质、特色突破、品牌辐射四项培养工程和分层关爱。',
        ],
        'changes': [
            '明确“四级”是专业发展阶段，不是行政等级、身份等级或薪酬等级。',
            '增加从8月至次年7月的年度运行周期：确定目标、实践改进、中期面谈、成果展示、年度评价。',
            '建立师德与育人、课堂与教学、教研与课程、团队与服务、发展增值五个评价领域。',
            '增加课堂观察、学生作品、育人案例、课程资源、带教效果和团队贡献等代表性证据。',
            '明确文化课、艺体综合、班主任、行政兼课和教辅岗位不能使用完全相同的一把尺子。',
            '评价形成“优势、改进方向、支持建议、下一年度目标”，不公开教师个人分数排名。',
            '增加教师复核、权益保护、困难帮扶和新增工作负担审核。',
            '晋级原则上每两至三年一次，由教师自主申报，不强制所有教师逐级竞争。',
        ],
        'summary': '“四维四级”是本轮保留最完整的核心构想，主要补足了年度怎么运行、怎么评价、怎么支持教师，以及怎样避免把专业发展变成刚性等级管理。',
    },
    {
        'title': '《“三红三专”一体化课程体系实施方案》',
        'keep': [
            '保留“三红铸魂＋三专赋能”两大课程集群。',
            '保留国家课程、校本必修、拓展选修、社团实践四类课程形态。',
            '保留七年级筑基、八年级发展、九年级圆梦三年递进。',
            '保留校史、本土红色文化、学习方法、心理健康、劳动、体育、美育、科创、无人机和人工智能等课程内容。',
        ],
        'changes': [
            '明确国家课程必须开齐开足，校本课程不能挤占国家课程课时。',
            '为四类课程分别明确对象、课时、责任边界和建设要求。',
            '增加课程建设优先级，2026—2027先做基础和急需课程，之后逐步发展拓展、跨学科和共享课程。',
            '建立课程准入制度，开课前必须说明对象、目标、课时、教师、场地、安全和评价。',
            '涉及校外机构、收费、人工智能平台、个人信息、实验和户外活动的，增加专项审核。',
            '增加单课统一操作结构：情境与目标—学习与实践—表达与反馈—行动与延伸。',
            '增加课程复盘和退出机制，形成保留、调整、扩大、暂停、停止五类结论。',
        ],
        'summary': '课程体系的大架构和特色内容都保留了，主要补上了哪些先做、谁来做、占什么课时、怎样开课、怎样评价以及效果不好怎样退出。',
    },
    {
        'title': '《集体备课规范化管理制度》',
        'keep': [
            '保留每周集体备课。',
            '保留上周学情复盘、下周主备说课、练习作业研讨、资源共享与任务分工四个核心环节。',
        ],
        'changes': [
            '重新界定“四统一”：统一课程进度、核心目标、重难点、质量底线和作业总量，不要求所有教师使用完全相同的课堂流程、教案和课件。',
            '允许教师根据不同班情进行二次设计。',
            '每次备课聚焦3—5个关键问题，不追求议程和材料数量。',
            '每次只形成“一张记录表＋一个共享资源包”，不再重复打印主备稿、说课稿、教案、课件、作业和照片。',
            '明确常规备课60—90分钟，小规模学科可隔周或跨年级联合。',
            '分别增加新授课、复习课、试卷讲评、实验实践课和毕业年级的备课重点。',
            '管理部门由周周全覆盖检查调整为抽样观察，重点看是否解决真实课堂和学生问题。',
            '把学生受益和教师负担同时纳入质量评价。',
        ],
        'summary': '集体备课的严格要求没有降低，但把“统一教案、统一课堂”调整为“统一质量底线、允许教师二次设计”，同时大幅减少重复材料。',
    },
    {
        'title': '《集体备课活动记录表》',
        'keep': [
            '保留基本信息、上周学情复盘、下周教学研讨、作业和练习确定、任务分工、备课组长和主备教师签字。',
        ],
        'changes': [
            '重新制作成横向、可直接填写的结构化表格。',
            '复盘部分只填写关键事实或数据、共同判断、补救安排和责任人，不再抄写大段教学内容。',
            '增加“允许教师依据班情调整的内容”，与二次设计制度相衔接。',
            '把作业分成课堂练习、课后基础作业、选择提升任务、个别支持任务、小测或周练，并记录预计时长、题量和反馈方式。',
            '增加责任人、完成时间、共享位置和下一次备课需反馈的问题。',
            '增加“实施后简要反馈”，形成真正闭环，而不是开会结束即归档。',
        ],
        'summary': '原表的主要栏目保留了，但改成“一张表完成复盘、决策、分工和后续反馈”，老师不用再重复抄写教案内容。',
    },
    {
        'title': '《初三A层次班级16周学业指导与生涯规划专题课教学提纲》',
        'keep': [
            '保留学习态度、学习能力、学习方法、学习策略四维框架。',
            '保留面向基础较好学生的16周课程。',
            '保留学业突破、心态调节、学习方法、考试策略和生涯规划等主题。',
        ],
        'changes': [
            '明确A层是阶段性学习支持分组，不是固定身份，也不公开给学生贴“尖子生”标签。',
            '进入和调整不只依据一次考试排名，还结合近期学习需要、本人意愿和教师专业判断。',
            '删除“满分公式”“顶尖公式”等过度承诺，将目标调整为稳定基础、突破瓶颈和可持续发展。',
            '每一周增加明确课堂任务和学生成果，如目标卡、得分结构图、错因图谱、稳态应对卡、生涯探索表和成长陈述。',
            '文科、理科、英语等专题由相应学科骨干共同承担，避免通用方法替代学科专业教学。',
            '增加基线、期中和期末三次同一量表诊断。',
            '明确持续焦虑、压力和情绪困扰需要转介专业支持。',
            '不形成公开总分排名，期末主要提供个人四维成长反馈。',
        ],
        'summary': 'A层课程16周主体内容基本保留，重点降低标签化和过度应试色彩，让课程既能拔尖，也能兼顾心态、生涯和持续发展。',
    },
    {
        'title': '《初三C层次班级16周学业指导与生涯规划专题课教学提纲》',
        'keep': [
            '保留对基础较弱学生的理解和共情。',
            '保留学习态度、能力、方法、策略四维框架。',
            '保留先稳心态、再养习惯、再提能力，以及基础优先、低压力、小步进步的核心设计。',
            '保留16周课程和多元生涯指导。',
        ],
        'changes': [
            '明确C层是阶段性支持分组，不在公开场合使用“差生”“后进生”等称呼。',
            '保留理解学生委屈、承认学生努力的开篇思想，但压缩为更适合课堂的核心表达。',
            '第一课增加匿名表达和保密约定，让学生能够说出希望老师理解的问题。',
            '每周配置一个低门槛成果，如极简复盘卡、基础清单、情绪与求助卡、一周三任务计划、路径信息卡和支持网络图。',
            '建立普惠课堂支持、导师短时支持、家校协同、专业转介四层支持。',
            '对持续低落、自伤风险和严重家庭冲突增加专业报告与转介要求。',
            '关于普高、职业教育等路径，只提供真实信息，不作保证性承诺。',
            '期末不使用固定等级和公开排名，而采用“明显进步、稳定改善、开始行动、仍需支持”等描述。',
        ],
        'summary': 'C层课程最核心的共情和兜底育人思想完整保留，主要进一步保护学生尊严，增加导师、家校和专业支持，使课程不只是一次讲话，而是一套连续支持机制。',
    },
    {
        'title': '《全员素养提升暨创新拔尖培育校本课程提纲》',
        'keep': [
            '保留创新潜质发展课程和全员基础素养课程两大模块。',
            '保留全年32课时。',
            '保留创新思维、科创实践、品格、习惯、学习方法、身心健康和生涯教育。',
            '保留党委书记领航课程的特色定位。',
        ],
        'changes': [
            '将“党委书记全程亲授全部课程”优化为“党委书记领航关键主题，学科骨干、心理教师、生涯教师、科创导师和班主任团队协同实施”。',
            '不再固定按照“创新班、普通班”判断学生价值，改为创新潜质发展组和基础素养支持组，并允许动态调整。',
            '不以一次考试作为进入创新模块的唯一依据。',
            '为两个模块建立共同基础课程：学校文化、规则自律、学习方法、身心健康、责任合作、数字与人工智能素养和生涯启蒙。',
            '每一课增加具体学生成果，避免只写授课主题。',
            '增加人工智能、实验、无人机、校外实践和个人信息的安全边界。',
            '明确每学期对学生获得感、课程负担和实际成果进行复盘，避免32节课变成单向讲座。',
        ],
        'summary': '“书记领航、全员提升、创新拔尖”的核心没有改变，主要把一人授课优化为书记领航、专业团队共同实施，并增加动态选择和课程成果要求。',
    },
    {
        'title': '《“十五五”（2026—2030）发展规划建议》',
        'keep': [
            '保留红专文化、三红三专育人、扩优提质、健康第一、强师兴教、数字赋能、课程课堂改革、学生评价、教师四维四级、家校社协同、安全心理和质量提升等主要构想。',
        ],
        'changes': [
            '明确文件身份是党委书记办学构想和正式五年规划编制的参考建议稿，不与学校正式规划并行。',
            '不使用未经核实的生源比例、教师结构和精确提升数据，提出2026年先核定真实基线。',
            '增加学校发展基础、现实问题和2030年应达到的变化。',
            '形成2026—2027整改筑基、2028—2029协同提质、2030模式定型三个阶段。',
            '把十星、标杆班级、四维四级、分层课程等构想纳入12项学校重点工程，而不是增加新的上位体系。',
            '明确每年只从12项工程中选择6—8项重点推进，不要求全部项目同时铺开。',
            '增加项目任务书、中期协调、学期复盘、年度报告和规划调整机制。',
            '指标只提出方向和目标类型，真实基线和具体比例在2026年核定后确定。',
            '增加2026统一与整改、2027机制筑基、2028专业提质、2029协同扩优、2030评估定型的逐年安排。',
        ],
        'summary': '原有战略方向没有改，主要把宏大的五年设想转化为有基线、有阶段、有重点工程、有年度选择和有退出调整机制的规划建议。',
    },
]


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_east_asia(run, font):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('— ')
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    run2 = paragraph.add_run(' —')
    for r in (run, run2):
        set_east_asia(r, '宋体'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(102,102,102)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = instruction
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = '目录将在Word中更新后显示'
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar, instrText, fldChar2, text, fldChar3])


def add_paragraph(doc, text='', style=None, bold=False, color=None, align=None, size=None, font='宋体', indent=True, space_after=5):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_east_asia(r, font)
    if size: r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(item)
        set_east_asia(r, '宋体'); r.font.size = Pt(10.5)


def add_callout(doc, title, text, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0,0)
    cell.width = Cm(16.3)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + '：')
    set_east_asia(r, '黑体'); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor.from_string(RED)
    r2 = p.add_run(text)
    set_east_asia(r2, '宋体'); r2.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        set_cell_shading(hdr[i], NAVY)
        set_cell_margins(hdr[i], 120, 110, 120, 110)
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_east_asia(r, '黑体'); r.font.size = Pt(font_size); r.bold = True; r.font.color.rgb = RGBColor.from_string(WHITE)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            set_cell_margins(cells[i], 100, 100, 100, 100)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], 'F7F9FB')
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(str(val)); set_east_asia(r, '宋体'); r.font.size = Pt(font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = '宋体'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.45
    for name, size, color, font, before, after in [
        ('Title', 26, RED, '方正小标宋简体', 0, 16),
        ('Heading 1', 16, RED, '黑体', 16, 8),
        ('Heading 2', 13, NAVY, '黑体', 10, 5),
        ('Heading 3', 11, NAVY, '黑体', 7, 3),
    ]:
        s = styles[name]
        s.font.name = font; s._element.rPr.rFonts.set(qn('w:eastAsia'), font); s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color); s.font.bold = True
        s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    if 'SubtitleCN' not in styles:
        s = styles.add_style('SubtitleCN', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = '楷体_GB2312'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'楷体_GB2312'); s.font.size = Pt(15); s.font.color.rgb = RGBColor.from_string(NAVY)
        s.paragraph_format.space_after = Pt(12)


def main():
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.2); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.3)
    sec.header_distance = Cm(1.1); sec.footer_distance = Cm(1.1)

    # header/footer
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run('开江中学实验学校党委书记15份参考文件优化完善说明')
    set_east_asia(hr,'宋体'); hr.font.size = Pt(9); hr.font.color.rgb = RGBColor(110,110,110)
    add_page_number(sec.footer.paragraphs[0])

    # cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(78); p.paragraph_format.space_after = Pt(18)
    r = p.add_run('开江中学实验学校')
    set_east_asia(r,'黑体'); r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(NAVY); r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(24)
    r = p.add_run('党委书记15份参考文件')
    set_east_asia(r,'方正小标宋简体'); r.font.size = Pt(27); r.font.color.rgb = RGBColor.from_string(RED); r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12)
    r = p.add_run('优化完善说明')
    set_east_asia(r,'方正小标宋简体'); r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(RED); r.bold = True
    p = doc.add_paragraph(style='SubtitleCN'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('——原稿主体保留、操作机制补强、风险边界完善').font.size = Pt(15)
    add_callout(doc, '文件用途', '本说明用于向学校党委书记直观汇报15份原始文件分别保留了什么、主要优化了什么，以及优化后的价值变化。', LIGHT_GOLD)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(95); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('内部沟通稿')
    set_east_asia(r,'黑体'); r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('2026年7月')
    set_east_asia(r,'宋体'); r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()

    # Preface
    doc.add_heading('说明', level=1)
    add_paragraph(doc, '本轮优化坚持“保留主体、修正问题、补足操作、控制风险”的基本原则。没有改变学校党委书记提出的“红专十星、十大标杆班级、三红三专课程、教师四维四级、分层学业指导和十五五发展”等核心构想，而是根据每份文件的功能和实际使用场景，进行有限度、针对性的完善。')
    add_callout(doc, '四项主要转换', '从办学构想转化为可执行流程；从多套积分和大量留痕转化为关键事实与代表性证据；从简单扣分、排名和一票否决转化为发展性评价与规范程序；从一次全面铺开转化为试行、复盘、完善、推广。', LIGHT_RED)

    doc.add_heading('一、总体优化概览', level=1)
    overview_rows = []
    for i, f in enumerate(FILES,1):
        overview_rows.append([str(i), f['title'].replace('《','').replace('》',''), f['summary']])
    add_table(doc, ['序号','文件名称','优化后的核心变化'], overview_rows, widths=[1.2,6.0,9.0], font_size=8.6)

    doc.add_heading('二、统一优化原则', level=1)
    add_table(doc, ['优化维度','主要处理'], [
        ['主体保留','原稿名称、核心项目、主要结构和办学意图原则上保留，不另起炉灶。'],
        ['实操补足','增加对象、流程、时间、责任、材料、评价、复核和配套表单。'],
        ['减负增效','压减重复积分、照片台账、长篇材料和多头报送，提倡一次采集、多处使用。'],
        ['发展评价','降低单纯排名、扣分和一票否决，增加个人增值、整改进步和代表性事实。'],
        ['权益保护','完善申诉复核、隐私保护、教师和学生表达、专业转介及结果使用边界。'],
        ['循序渐进','明确试行范围、阶段安排、复盘修订和退出机制，不要求所有项目一次铺开。'],
    ], widths=[3.4,12.6])

    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('目  录'); set_east_asia(r,'方正小标宋简体'); r.font.size = Pt(22); r.bold = True; r.font.color.rgb = RGBColor.from_string(RED)
    toc = doc.add_paragraph(); add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    # details
    doc.add_heading('三、15份文件逐项优化说明', level=1)
    for idx, f in enumerate(FILES,1):
        doc.add_heading(f'{idx}. {f["title"]}', level=2)
        doc.add_heading('（一）保留的主体内容', level=3)
        add_bullets(doc, f['keep'])
        doc.add_heading('（二）主要优化内容', level=3)
        for n, item in enumerate(f['changes'],1):
            add_paragraph(doc, f'{n}. {item}', indent=False, space_after=3)
        add_callout(doc, '与书记沟通时可概括为', f['summary'], LIGHT_BLUE)
        if idx in [5,8,11,15] and idx != 15:
            doc.add_page_break()

    doc.add_heading('四、排版和文档使用方面的统一优化', level=1)
    add_bullets(doc, [
        '统一正式封面、文件名称、标题层级、正文格式、页眉页脚和表格样式。',
        '重要流程、标准、责任和阶段安排尽量改为结构化表格，便于阅读、审议和执行。',
        '实施类文件增加可直接填写的申报表、评分表、任务书、反馈表和记录表。',
        '删除“AI生成”“可信度10分”等不适合作为正式学校文件的说明。',
        '删除文件末尾不必要的参考来源附带信息，正文只保留与执行直接相关的内容。',
        '全部形成可独立编辑、审议和试行的Word文件，并完成转换和页面检查。',
    ])

    doc.add_heading('五、与党委书记交流的建议表述', level=1)
    add_callout(doc, '建议总说明', '这次优化没有改变您提出的“红专十星、十大标杆班级、三红三专课程、教师四维四级、分层学业指导和十五五发展”等主体构想。主要做了四方面完善：一是修正少量重复、分值冲突和前后口径不一致；二是补充责任主体、时间节点、实施流程和操作表格；三是适当降低重复积分、材料留痕和教师负担；四是完善学生教师权益、申诉复核、隐私保护和分阶段试行机制。总体上是保留您的办学思想，让这些构想更加规范、稳妥，也更便于学校教师直接执行。', LIGHT_RED)

    doc.add_heading('六、总体结论', level=1)
    add_paragraph(doc, '15份原始文件所体现的办学方向和制度创意具有较高保留价值。本轮优化不是替代书记构想，而是通过程序化、标准化、减负化和风险边界完善，使这些构想从“可以讨论的设计稿”进一步转化为“可以审议、可以试行、可以复盘完善的学校文件”。')
    add_paragraph(doc, '后续正式发布时，仍建议由学校核定处室正式名称、人员分工、年度时间、真实基线数据和本地政策要求；涉及学生处分、教师绩效职称、心理健康、安全、数据和收费的事项，应按相应程序完成专项审核。')

    # core properties
    doc.core_properties.title = '开江中学实验学校党委书记15份参考文件优化完善说明'
    doc.core_properties.subject = '15份办学构想和制度参考稿优化内容说明'
    doc.core_properties.author = '开江中学实验学校文件优化工作'
    doc.save(DOCX)
    print(DOCX)

if __name__ == '__main__':
    main()
