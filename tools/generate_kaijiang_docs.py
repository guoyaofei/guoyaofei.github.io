from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List
import os
import zipfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from kaijiang_docs_content import DOCS

OUT_DIR = Path(os.environ.get("OUT_DIR", "generated_docs"))
OUT_DIR.mkdir(parents=True, exist_ok=True)

THEME_RED = "9E1B32"
THEME_NAVY = "1F4E79"
THEME_GOLD = "C79A3B"
THEME_LIGHT = "F4F6F8"
THEME_RED_LIGHT = "FBECEF"
THEME_BLUE_LIGHT = "EAF2F8"
TEXT_DARK = "222222"
TEXT_GREY = "666666"


def set_east_asia_font(run, font_name: str, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:%s" % edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in kwargs[edge]:
                    element.set(qn("w:%s" % key), str(kwargs[edge][key]))


def set_table_borders(table, color: str = "B7C3D0", size: int = 6):
    for row in table.rows:
        for cell in row.cells:
            border = {"val": "single", "sz": size, "color": color, "space": "0"}
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_east_asia_font(run, "宋体", 9, color=TEXT_GREY)


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_east_asia_font(r, "黑体", 18, True, THEME_NAVY)
    p.paragraph_format.space_after = Pt(12)
    p2 = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p2._p.append(fld)
    doc.add_page_break()


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(0.85)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, font_name, size, color, before, after in [
        ("Title", "方正小标宋简体", 26, THEME_RED, 0, 18),
        ("Subtitle", "楷体_GB2312", 15, THEME_NAVY, 0, 10),
        ("Heading 1", "黑体", 17, THEME_RED, 18, 10),
        ("Heading 2", "黑体", 15, THEME_NAVY, 14, 7),
        ("Heading 3", "楷体_GB2312", 13.5, TEXT_DARK, 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = font_name
        st._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Callout" not in styles:
        st = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles["Callout"]
    st.font.name = "楷体_GB2312"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体_GB2312")
    st.font.size = Pt(12.5)
    st.font.color.rgb = RGBColor.from_string(THEME_RED)
    st.paragraph_format.left_indent = Cm(0.8)
    st.paragraph_format.right_indent = Cm(0.8)
    st.paragraph_format.space_before = Pt(6)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.35


def configure_sections(doc: Document, short_title: str):
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.55)
        section.right_margin = Cm(2.35)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(f"开江中学实验学校“红专”学校发展规划文件体系｜{short_title}")
        set_east_asia_font(hr, "宋体", 9, color=TEXT_GREY)
        fp = section.footer.paragraphs[0]
        add_page_number(fp)


def add_cover(doc: Document, meta: Dict[str, Any]):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(meta["title"])
    set_east_asia_font(r, "方正小标宋简体", 27, True, THEME_RED)

    if meta.get("subtitle"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(28)
        r = p.add_run(meta["subtitle"])
        set_east_asia_font(r, "楷体_GB2312", 16, True, THEME_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━")
    set_east_asia_font(r, "宋体", 12, color=THEME_GOLD)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("status", "校级审议稿"))
    set_east_asia_font(r, "黑体", 13, True, THEME_NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("开江中学实验学校")
    set_east_asia_font(r, "黑体", 15, True, TEXT_DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("date", "二〇二六年七月"))
    set_east_asia_font(r, "宋体", 12, color=TEXT_GREY)
    doc.add_page_break()


def add_overview(doc: Document, overview: List[List[str]]):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("编制说明")
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in overview:
        row = table.add_row()
        row.cells[0].width = Cm(3.1)
        row.cells[1].width = Cm(13.4)
        set_cell_shading(row.cells[0], THEME_RED_LIGHT)
        for idx, text in enumerate([label, value]):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.line_spacing = 1.3
            run = para.add_run(text)
            set_east_asia_font(run, "黑体" if idx == 0 else "宋体", 11.5, idx == 0, THEME_RED if idx == 0 else TEXT_DARK)
    set_table_borders(table)
    doc.add_paragraph()


def add_callout(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, THEME_RED_LIGHT)
    set_cell_border(cell,
                    top={"val": "single", "sz": 14, "color": THEME_RED, "space": "0"},
                    bottom={"val": "single", "sz": 6, "color": "E6BCC4", "space": "0"},
                    left={"val": "single", "sz": 14, "color": THEME_RED, "space": "0"},
                    right={"val": "single", "sz": 6, "color": "E6BCC4", "space": "0"})
    para = cell.paragraphs[0]
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0.3)
    para.paragraph_format.right_indent = Cm(0.3)
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(text)
    set_east_asia_font(run, "楷体_GB2312", 12.5, True, THEME_RED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_east_asia_font(r1, "黑体", 12.5, True, TEXT_DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_east_asia_font(r2, "宋体", 12.5, False, TEXT_DARK)
    else:
        r = p.add_run(text)
        set_east_asia_font(r, "宋体", 12.5, False, TEXT_DARK)
    return p


def add_bullets(doc: Document, items: List[str], numbered: bool = False, level: int = 0):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7 + level * 0.55)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = 1.35
        marker = f"{i}. " if numbered else "• "
        r1 = p.add_run(marker)
        set_east_asia_font(r1, "黑体", 12, True, THEME_RED)
        r2 = p.add_run(item)
        set_east_asia_font(r2, "宋体", 12.3, False, TEXT_DARK)


def add_table(doc: Document, headers: List[str], rows: List[List[str]], widths: List[float] | None = None, font_size: float = 10.5, landscape: bool = False):
    previous_orientation = None
    if landscape:
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        previous_orientation = sec.orientation
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.6)
        sec.right_margin = Cm(1.6)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        set_cell_shading(cell, THEME_NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(h)
        set_east_asia_font(run, "黑体", font_size, True, "FFFFFF")
        if widths and j < len(widths):
            cell.width = Cm(widths[j])
    for i, row_data in enumerate(rows):
        row = table.add_row()
        fill = "FFFFFF" if i % 2 == 0 else THEME_LIGHT
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            if j == 0 and len(headers) <= 5:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(val))
            set_east_asia_font(run, "宋体", font_size, False, TEXT_DARK)
            if widths and j < len(widths):
                cell.width = Cm(widths[j])
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if landscape:
        sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
        sec2.orientation = WD_ORIENT.PORTRAIT
        sec2.page_width = Cm(21.0)
        sec2.page_height = Cm(29.7)
        sec2.top_margin = Cm(2.4)
        sec2.bottom_margin = Cm(2.2)
        sec2.left_margin = Cm(2.55)
        sec2.right_margin = Cm(2.35)


def add_signature(doc: Document, text: str = "本文件经学校规定程序审议通过后施行。"):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_east_asia_font(r, "楷体_GB2312", 11.5, False, TEXT_GREY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("开江中学实验学校")
    set_east_asia_font(r, "宋体", 11.5, False, TEXT_DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("二〇二六年七月")
    set_east_asia_font(r, "宋体", 11.5, False, TEXT_DARK)


def render_doc(meta: Dict[str, Any]) -> Path:
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    configure_sections(doc, meta.get("short_title", meta["title"]))
    add_cover(doc, meta)
    add_overview(doc, meta.get("overview", []))
    if meta.get("toc", True):
        add_toc(doc)

    for block in meta.get("blocks", []):
        typ = block["type"]
        if typ == "h1":
            doc.add_paragraph(block["text"], style="Heading 1")
        elif typ == "h2":
            doc.add_paragraph(block["text"], style="Heading 2")
        elif typ == "h3":
            doc.add_paragraph(block["text"], style="Heading 3")
        elif typ == "p":
            add_paragraph(doc, block["text"], block.get("bold_prefix"))
        elif typ == "callout":
            add_callout(doc, block["text"])
        elif typ == "bullets":
            add_bullets(doc, block["items"], block.get("numbered", False), block.get("level", 0))
        elif typ == "table":
            add_table(doc, block["headers"], block["rows"], block.get("widths"), block.get("font_size", 10.5), block.get("landscape", False))
        elif typ == "pagebreak":
            doc.add_page_break()
        elif typ == "signature":
            add_signature(doc, block.get("text", "本文件经学校规定程序审议通过后施行。"))
        else:
            raise ValueError(f"Unknown block type: {typ}")

    configure_sections(doc, meta.get("short_title", meta["title"]))
    out = OUT_DIR / meta["filename"]
    doc.save(out)
    return out


def main():
    generated = []
    for meta in DOCS:
        generated.append(render_doc(meta))
    readme = OUT_DIR / "00_文件清单与使用说明.txt"
    readme.write_text(
        "开江中学实验学校“红专”学校发展规划首批十份核心文件\n\n"
        + "\n".join(f"{i+1:02d}. {p.name}" for i, p in enumerate(generated))
        + "\n\n说明：本批文件均为校级审议稿，核心架构、术语、职责、指标及方案应经学校规定程序审议后发布。涉及具体人员、机构名称、基线数据、年度目标和政策依据的内容，发布前应以学校现行情况和正式文件复核。\n",
        encoding="utf-8",
    )
    zip_path = OUT_DIR / "开江中学实验学校红专学校发展规划首批十份核心文件.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in generated + [readme]:
            zf.write(p, p.name)
    print(f"Generated {len(generated)} documents and {zip_path}")


if __name__ == "__main__":
    main()
